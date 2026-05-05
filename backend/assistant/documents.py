"""Document ingestion for Zilo Chat.

- Accepts uploads (PDF, DOCX, TXT/MD, images).
- Extracts plain text where possible (pypdf / python-docx / raw decode).
- For images + PDFs, keeps the original base64 so Claude can receive them
  natively via its Messages API.
- Stores each document under `assistant_documents` keyed by conversation_id.
"""
from __future__ import annotations

import base64
import io
import logging
import uuid
from datetime import datetime
from typing import Any, Dict, List, Optional, Tuple

from .embeddings import (
    LONG_DOC_THRESHOLD,
    delete_document_chunks,
    ingest_document_chunks,
)

logger = logging.getLogger(__name__)

MAX_BYTES = 15 * 1024 * 1024          # 15 MB per file
MAX_EXTRACT_CHARS = 200_000           # cap stored text per file

SUPPORTED_MIME = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "text/plain": "txt",
    "text/markdown": "md",
    "text/csv": "csv",
    "image/png": "image",
    "image/jpeg": "image",
    "image/webp": "image",
    "image/gif": "image",
}


def _extract_pdf(data: bytes) -> str:
    text = ""
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(data))
        out: List[str] = []
        for page in reader.pages:
            try:
                out.append(page.extract_text() or "")
            except Exception:
                continue
        text = "\n\n".join(out).strip()
    except Exception as e:
        logger.warning(f"[documents] PDF extract failed: {e}")

    # If pypdf found no text (scanned/image-only PDF), fall back to AWS Textract
    if not text:
        text = _ocr_pdf_textract(data)
    return text


def _ocr_pdf_textract(data: bytes) -> str:
    """Use AWS Textract to extract text from a scanned PDF. Returns '' on any failure."""
    import os
    access_key = os.environ.get("AWS_ACCESS_KEY_ID", "").strip()
    secret_key = os.environ.get("AWS_SECRET_ACCESS_KEY", "").strip()
    region = os.environ.get("AWS_DEFAULT_REGION") or os.environ.get("AWS_REGION") or "us-east-1"
    if not access_key or not secret_key:
        logger.info("[documents] Textract OCR skipped — AWS credentials not configured.")
        return ""
    try:
        import boto3
        client = boto3.client(
            "textract",
            aws_access_key_id=access_key,
            aws_secret_access_key=secret_key,
            region_name=region,
        )
        # detect_document_text works on a single page at a time (bytes).
        # For multi-page PDFs we use start_document_text_detection (async),
        # but to keep it simple here we use the sync API (supports 1 page or single image).
        # Multi-page PDFs need S3 upload; we cap at 5 MB to use sync path.
        if len(data) <= 5 * 1024 * 1024:
            resp = client.detect_document_text(Document={"Bytes": data})
            lines: List[str] = [
                blk["Text"]
                for blk in (resp.get("Blocks") or [])
                if blk.get("BlockType") == "LINE"
            ]
            return "\n".join(lines).strip()
        # Larger PDFs: skip OCR and inform caller
        logger.info("[documents] PDF too large for sync Textract (>5 MB). OCR skipped.")
        return ""
    except Exception as e:
        logger.warning(f"[documents] Textract OCR failed: {e}")
        return ""


def _extract_docx(data: bytes) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(data))
        parts: List[str] = [p.text for p in doc.paragraphs if p.text]
        for t in doc.tables:
            for row in t.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        return "\n".join(parts).strip()
    except Exception as e:
        logger.warning(f"[documents] DOCX extract failed: {e}")
        return ""


def _extract_text(data: bytes) -> str:
    for enc in ("utf-8", "utf-16", "latin-1"):
        try:
            return data.decode(enc)
        except Exception:
            continue
    return ""


def extract(content: bytes, mime_type: str) -> Tuple[str, str]:
    """Return (kind, extracted_text). Kind is one of: pdf, docx, txt, md, csv, image."""
    kind = SUPPORTED_MIME.get(mime_type, "")
    if not kind:
        return "", ""
    if kind == "pdf":
        text = _extract_pdf(content)
    elif kind == "docx":
        text = _extract_docx(content)
    elif kind in ("txt", "md", "csv"):
        text = _extract_text(content)
    else:
        text = ""  # images: no text extraction, Claude handles them natively
    return kind, text[:MAX_EXTRACT_CHARS]


async def _upload_image_to_s3(content: bytes, filename: str, mime_type: str) -> Optional[str]:
    """Upload image bytes to S3/Cloudinary and return a public URL. Returns None on failure."""
    try:
        b64_str = base64.b64encode(content).decode("ascii")
        from image_handler import S3Handler
        url = await S3Handler.upload_file(b64_str, filename, content_type=mime_type)
        return url or None
    except Exception as e:
        logger.warning("[documents] S3 image upload failed: %s", e)
    return None


async def store_upload(
    db,
    *,
    user_id: str,
    conversation_id: str,
    filename: str,
    mime_type: str,
    content: bytes,
) -> Dict[str, Any]:
    if len(content) > MAX_BYTES:
        raise ValueError(f"File too large. Max is {MAX_BYTES // (1024*1024)} MB.")
    if mime_type not in SUPPORTED_MIME:
        raise ValueError(
            f"Unsupported file type '{mime_type}'. "
            "Supported: PDF, DOCX, TXT, MD, CSV, PNG, JPEG, WEBP, GIF."
        )
    kind, text = extract(content, mime_type)

    # For images: upload to S3 so design tools can use the URL as product_image_url
    public_url: Optional[str] = None
    if kind == "image":
        public_url = await _upload_image_to_s3(content, filename, mime_type)

    doc_id = str(uuid.uuid4())
    doc = {
        "_id": doc_id,
        "user_id": user_id,
        "conversation_id": conversation_id,
        "filename": filename,
        "mime_type": mime_type,
        "kind": kind,
        "size": len(content),
        "text": text,
        "text_len": len(text),
        # base64 is only kept for images + PDFs so Claude can ingest them natively
        "b64": base64.b64encode(content).decode("ascii") if kind in ("image", "pdf") else None,
        # public S3 URL for images — passed to design tools as product_image_url
        "public_url": public_url,
        "created_at": datetime.utcnow(),
    }
    await db.assistant_documents.insert_one(doc)

    # Chunk + embed long text docs so the agent can retrieve excerpts via
    # `search_documents` instead of drowning in a giant preamble.
    chunks_indexed = 0
    if text and len(text) >= LONG_DOC_THRESHOLD:
        try:
            chunks_indexed = await ingest_document_chunks(
                db,
                user_id=user_id,
                conversation_id=conversation_id,
                document_id=doc_id,
                filename=filename,
                text=text,
            )
        except Exception as e:
            logger.warning(f"[documents] embedding ingest failed: {e}")

    return {
        "id": doc_id,
        "filename": filename,
        "kind": kind,
        "mime_type": mime_type,
        "size": len(content),
        "text_len": len(text),
        "has_text": bool(text),
        "chunks_indexed": chunks_indexed,
        "public_url": public_url,
    }


async def list_for_conversation(db, user_id: str, conversation_id: str) -> List[Dict[str, Any]]:
    rows = await db.assistant_documents.find(
        {"user_id": user_id, "conversation_id": conversation_id}
    ).sort("created_at", 1).to_list(50)
    return [{
        "id": r["_id"],
        "filename": r.get("filename"),
        "kind": r.get("kind"),
        "mime_type": r.get("mime_type"),
        "size": r.get("size"),
        "text_len": r.get("text_len", 0),
        "has_text": bool(r.get("text")),
        "created_at": r.get("created_at"),
        "public_url": r.get("public_url"),
    } for r in rows]


async def load_full(db, user_id: str, conversation_id: str) -> List[Dict[str, Any]]:
    """Load docs with their text + base64 payloads for passing into the model."""
    return await db.assistant_documents.find(
        {"user_id": user_id, "conversation_id": conversation_id}
    ).sort("created_at", 1).to_list(50)


async def delete_document(db, user_id: str, doc_id: str) -> bool:
    res = await db.assistant_documents.delete_one({"_id": doc_id, "user_id": user_id})
    # Best-effort cleanup of the chunk index
    try:
        await delete_document_chunks(db, user_id, doc_id)
    except Exception:
        pass
    return res.deleted_count > 0


def build_context_preamble(docs: List[Dict[str, Any]], per_doc_chars: int = 12_000) -> Optional[str]:
    """Return a single string to prepend as a system message so any provider
    (OpenAI / DeepSeek / Grok) can reference the documents. For Claude we ALSO
    pass native image/PDF blocks on top."""
    if not docs:
        return None
    blocks: List[str] = ["The user has attached the following reference documents:"]
    image_urls: List[str] = []
    for i, d in enumerate(docs, 1):
        fn = d.get("filename") or f"document {i}"
        kind = d.get("kind") or "file"
        text = (d.get("text") or "").strip()
        public_url = d.get("public_url") or ""
        header = f"\n--- Document {i}: {fn} ({kind}) ---"
        if kind == "image" and public_url:
            image_urls.append(public_url)
            blocks.append(
                f"{header}\n"
                f"Public image URL (use this as product_image_url for design tools): {public_url}\n"
                f"[Image is also attached natively for visual reference]"
            )
        elif not text:
            blocks.append(header + "\n[No extractable text — image or scanned PDF]")
        else:
            snippet = text[:per_doc_chars]
            truncated = " …[truncated]" if len(text) > per_doc_chars else ""
            blocks.append(f"{header}\n{snippet}{truncated}")
    blocks.append(
        "\nWhen the user asks about these documents, quote short excerpts "
        "and cite them by filename. Do not fabricate content not present in the documents."
    )
    if image_urls:
        blocks.append(
            "\nDESIGN TOOL HINT: The user has uploaded "
            + ("an image" if len(image_urls) == 1 else f"{len(image_urls)} images")
            + " for design work. When generating social posts, ads, or any visual design, "
            "pass the image URL above as `product_image_url` to the design tools "
            "(generate_social_post, generate_ad_creative, generate_carousel_cover, refine_design, edit_product_image). "
            "Do NOT ask the user for an image URL — you already have it."
        )
    return "\n".join(blocks)
