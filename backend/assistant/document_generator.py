"""
document_generator.py
Converts Claude's markdown output into polished PDF and Word documents.
Input: plain markdown string (what the assistant already produces).
"""
from __future__ import annotations

import os
import re
import subprocess
import sys
import tempfile
import uuid
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

TEMP_DIR = Path(tempfile.gettempdir()) / "zilo_docs"
TEMP_DIR.mkdir(parents=True, exist_ok=True)

from .zilo_excel_styles import (
    title_band, section_hdr, col_hdr, label, value as zilo_value, total_row as zilo_total_row,
    gap as zilo_gap, set_col as zilo_set_col, setup_sheet as zilo_setup_sheet, tab_color as zilo_tab_color,
    DARK, GREEN, WHITE, LGRAY, MGRAY, AMBER_L, GREEN_L, GREEN_D, AMBER, RED_L, RED_D, DGRAY, BLUE, BLACK, GREEN_TX, MUTED, CARD
)

# Backend root (…/backend) — for local /uploads/ logo paths
ROOT_DIR = Path(__file__).resolve().parent.parent

# ── In-memory HTML preview store ─────────────────────────────────────────────
# Keyed by a UUID hex so the frontend can embed an <iframe> without S3/auth.
# Lives for the lifetime of the server process — good enough for a chat session.
_html_preview_store: Dict[str, str] = {}


def store_html_preview(html: str) -> str:
    """Cache an HTML document string and return a short retrieval key."""
    key = uuid.uuid4().hex
    _html_preview_store[key] = html
    return key


def get_html_preview(key: str) -> Optional[str]:
    return _html_preview_store.get(key)


def _backend_origin() -> str:
    port = (os.environ.get("PORT") or "8000").strip()
    base = (
        os.environ.get("BACKEND_PUBLIC_URL")
        or os.environ.get("WEBHOOK_BASE_URL")
        or os.environ.get("API_BASE_URL")
        or f"http://127.0.0.1:{port}"
    ).rstrip("/")
    return base


def _embed_image_as_data_uri(url: str) -> str:
    """Fetch an image and return a data: URI so iframe srcDoc and Playwright PDF can render it."""
    url = (url or "").strip()
    if not url:
        return ""
    if url.startswith("data:"):
        return url

    import base64
    import logging

    from image_handler import S3Handler

    resolved = S3Handler.resolve_accessible_url(url)
    if resolved.startswith("/api/"):
        resolved = f"{_backend_origin()}{resolved}"
    if resolved.startswith("/uploads/"):
        local = ROOT_DIR / resolved.lstrip("/").replace("/", os.sep)
        if local.is_file():
            import base64
            ext = local.suffix.lower().lstrip(".")
            ct = {"png": "image/png", "jpg": "image/jpeg", "jpeg": "image/jpeg", "webp": "image/webp", "gif": "image/gif"}.get(ext, "image/png")
            return f"data:{ct};base64,{base64.b64encode(local.read_bytes()).decode()}"
        resolved = f"{_backend_origin()}{resolved}"

    # Direct S3 fetch when credentials are available (avoids auth on internal routes)
    if S3Handler.parse_s3_source_to_bucket_key(resolved)[1] or "/api/images/s3/" in resolved:
        key = ""
        if "/api/images/s3/" in resolved:
            key = resolved.split("/api/images/s3/", 1)[1].split("?")[0].lstrip("/")
        else:
            _, key = S3Handler.parse_s3_source_to_bucket_key(resolved)
        bucket = (os.environ.get("AWS_BUCKET_NAME") or "").strip()
        if bucket and key:
            try:
                obj = S3Handler.get_s3_client().get_object(Bucket=bucket, Key=key)
                body = obj["Body"].read()
                ct = obj.get("ContentType") or "image/png"
                if ";" in ct:
                    ct = ct.split(";")[0].strip()
                return f"data:{ct};base64,{base64.b64encode(body).decode()}"
            except Exception as exc:
                logging.warning("[document_generator] S3 logo fetch failed: %s", exc)

    try:
        import httpx
        with httpx.Client(timeout=20.0, follow_redirects=True) as client:
            resp = client.get(resolved)
        resp.raise_for_status()
        ct = (resp.headers.get("content-type") or "image/png").split(";")[0].strip()
        return f"data:{ct};base64,{base64.b64encode(resp.content).decode()}"
    except Exception as exc:
        logging.warning("[document_generator] HTTP logo fetch failed (%s): %s", resolved[:80], exc)
        return resolved if resolved.startswith("http") else ""


# ─────────────────────────────────────────────────────────────────────────────
# Shared markdown parser  (uses stdlib `markdown`)
# ─────────────────────────────────────────────────────────────────────────────
def _fix_text_encoding(text: str) -> str:
    """Repair UTF-8 mojibake (e.g. â€" → –) from Windows-1252 / Latin-1 mis-decoding."""
    if not text:
        return text
    if "â" in text or "Ã" in text:
        # UTF-8 bytes were read as Windows-1252 (â + € + smart quote)
        try:
            repaired = text.encode("cp1252").decode("utf-8")
            if repaired and "\ufffd" not in repaired:
                text = repaired
        except (UnicodeEncodeError, UnicodeDecodeError):
            pass
        if "â" in text or "Ã" in text:
            try:
                repaired = text.encode("latin-1").decode("utf-8")
                if repaired and "\ufffd" not in repaired:
                    text = repaired
            except (UnicodeEncodeError, UnicodeDecodeError):
                pass
    for bad, good in (
        ("â€\u201c", "\u2013"),
        ("â€\u201d", "\u2014"),
        ("â€\x93", "\u2013"),
        ("â€\x94", "\u2014"),
        ("â€\u2018", "'"),
        ("â€\u2019", "'"),
        ("â€\x99", "'"),
        ("â€\x9c", "\u201c"),
        ("â€\x9d", "\u201d"),
        ("â€˜", "'"),
        ("Â ", " "),
        ("Â·", "\u00b7"),
    ):
        text = text.replace(bad, good)
    return text


def _ensure_blank_line_before_tables(md: str) -> str:
    """GitHub-style tables need a blank line before the first | row."""
    lines = md.split("\n")
    out: List[str] = []
    for line in lines:
        s = line.strip()
        if s.startswith("|") and out:
            prev = out[-1].strip()
            if prev and not prev.startswith("|") and not re.match(r"^[-|:|\s]+$", prev):
                out.append("")
        out.append(line)
    return "\n".join(out)


def _prepare_doc_markdown(md: str) -> str:
    md = _fix_text_encoding(md)
    return _ensure_blank_line_before_tables(md)


def _repair_html_paragraph_tables(html: str) -> str:
    """Convert | table | rows stuck inside <p>…</p> into real HTML tables."""
    import html as html_module

    def _repl(match: re.Match) -> str:
        inner = match.group(1)
        plain = re.sub(r"<br\s*/?>", "\n", inner)
        plain = re.sub(r"<[^>]+>", "", plain)
        lines = [ln.strip() for ln in plain.splitlines() if ln.strip().startswith("|")]
        if len(lines) < 2:
            return match.group(0)
        rows: List[List[str]] = []
        for ln in lines:
            compact = ln.replace(" ", "")
            if re.match(r"^\|[-:|]+\|$", compact):
                continue
            cells = [c.strip() for c in ln.strip().strip("|").split("|")]
            if cells:
                rows.append(cells)
        if not rows:
            return match.group(0)
        head = rows[0]
        body = rows[1:]
        thead = "<thead><tr>" + "".join(
            f"<th>{html_module.escape(c)}</th>" for c in head
        ) + "</tr></thead>"
        tbody = "".join(
            "<tr>" + "".join(f"<td>{html_module.escape(c)}</td>" for c in row) + "</tr>"
            for row in body
        )
        return f'<div class="table-wrap"><table>{thead}<tbody>{tbody}</tbody></table></div>'

    return re.sub(r"<p>((?:(?!</p>).)*\|(?:(?!</p>).)*)</p>", _repl, html, flags=re.DOTALL)


def _md_to_html(md: str) -> str:
    import markdown as _md
    md = _prepare_doc_markdown(md)
    html = _md.markdown(
        md,
        extensions=["tables", "fenced_code", "nl2br", "sane_lists"],
    )
    return _repair_html_paragraph_tables(html)


# ─────────────────────────────────────────────────────────────────────────────
# Block-level token walker
# ─────────────────────────────────────────────────────────────────────────────
_BLOCK_RE = re.compile(
    r"(<h[1-6][^>]*>.*?</h[1-6]>|"
    r"<pre><code[^>]*>.*?</code></pre>|"
    r"<table>.*?</table>|"
    r"<ul>.*?</ul>|"
    r"<ol>.*?</ol>|"
    r"<hr\s*/?>|"
    r"<p>.*?</p>)",
    re.DOTALL,
)


def _strip(html: str) -> str:
    return re.sub(r"<[^>]+>", "", html).strip()


def _parse_blocks(html: str) -> List[Tuple[str, str]]:
    """Return list of (type, raw_html) tuples."""
    blocks: List[Tuple[str, str]] = []
    for m in _BLOCK_RE.finditer(html):
        raw = m.group(0).strip()
        if raw.startswith("<h1"):
            blocks.append(("h1", raw))
        elif raw.startswith("<h2"):
            blocks.append(("h2", raw))
        elif raw.startswith("<h3"):
            blocks.append(("h3", raw))
        elif raw.startswith("<h4") or raw.startswith("<h5") or raw.startswith("<h6"):
            blocks.append(("h4", raw))
        elif raw.startswith("<pre"):
            code = re.sub(r"</?(?:pre|code)[^>]*>", "", raw).strip()
            blocks.append(("code", code))
        elif raw.startswith("<table"):
            blocks.append(("table", raw))
        elif raw.startswith("<ul"):
            items = re.findall(r"<li>(.*?)</li>", raw, re.DOTALL)
            for item in items:
                blocks.append(("ul", _strip(item)))
        elif raw.startswith("<ol"):
            items = re.findall(r"<li>(.*?)</li>", raw, re.DOTALL)
            for i, item in enumerate(items):
                blocks.append(("ol", f"{i + 1}. {_strip(item)}"))
        elif re.match(r"<hr", raw):
            blocks.append(("hr", ""))
        elif raw.startswith("<p"):
            text = _strip(raw)
            if text:
                blocks.append(("p", text))
    return blocks


def _parse_table(raw: str) -> List[List[str]]:
    rows: List[List[str]] = []
    for row in re.findall(r"<tr[^>]*>(.*?)</tr>", raw, re.DOTALL):
        cells = re.findall(r"<t[dh][^>]*>(.*?)</t[dh]>", row, re.DOTALL)
        rows.append([_strip(c) for c in cells])
    return rows


# ─────────────────────────────────────────────────────────────────────────────
# Style helpers
# ─────────────────────────────────────────────────────────────────────────────

def _resolve_fonts(font_style: str) -> tuple[str, str]:
    """Return (body_font, bold_font) ReportLab names for a style keyword."""
    style = (font_style or "").lower()
    if style in ("serif", "classic"):
        return "Times-Roman", "Times-Bold"
    # sans-serif / modern / minimal / default
    return "Helvetica", "Helvetica-Bold"


def _safe_hex(hex_str: str, fallback: str) -> str:
    """Return hex_str if it looks like a valid hex colour, else fallback."""
    h = (hex_str or "").strip()
    if re.match(r"^#[0-9A-Fa-f]{3,8}$", h):
        return h
    return fallback


# ─────────────────────────────────────────────────────────────────────────────
# PDF  (reportlab)
# ─────────────────────────────────────────────────────────────────────────────
def generate_pdf(
    markdown_content: str,
    filename: str | None = None,
    business_name: str = "",
    style: dict | None = None,
) -> str:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        HRFlowable,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    style = style or {}

    filename = filename or f"zilo_{uuid.uuid4().hex[:8]}.pdf"
    if not filename.endswith(".pdf"):
        filename += ".pdf"
    filepath = TEMP_DIR / filename

    doc = SimpleDocTemplate(
        str(filepath),
        pagesize=A4,
        rightMargin=2.2 * cm,
        leftMargin=2.2 * cm,
        topMargin=2.4 * cm,
        bottomMargin=2.2 * cm,
        title=filename.replace(".pdf", ""),
        author=business_name or "Zilo Chat",
    )

    # ── colour palette — use saved brand colors when available ───────────────
    PRIMARY_HEX   = _safe_hex(style.get("primary_color", ""), "#4F46E5")
    SECONDARY_HEX = _safe_hex(style.get("secondary_color", ""), "#EEF2FF")

    ACCENT      = colors.HexColor(PRIMARY_HEX)
    TABLE_HEAD  = colors.HexColor(SECONDARY_HEX)
    INK         = colors.HexColor("#111827")
    BODY        = colors.HexColor("#374151")
    MUTED       = colors.HexColor("#6B7280")
    RULE        = colors.HexColor("#E5E7EB")
    CODE_BG     = colors.HexColor("#F3F4F6")
    TABLE_ALT   = colors.HexColor("#F9FAFB")

    # ── fonts — serif vs sans-serif based on style profile ──────────────────
    BODY_FONT, BOLD_FONT = _resolve_fonts(style.get("font_style", ""))

    base = getSampleStyleSheet()

    def _s(**kw) -> ParagraphStyle:
        parent = kw.pop("parent", base["Normal"])
        name = kw.pop("name")
        return ParagraphStyle(name, parent=parent, **kw)

    H1   = _s(name="ZH1",   parent=base["Title"],   fontSize=22, leading=28, spaceAfter=4,  textColor=INK,   fontName=BOLD_FONT)
    H2   = _s(name="ZH2",   parent=base["Heading2"], fontSize=16, leading=20, spaceBefore=18, spaceAfter=4,  textColor=INK,   fontName=BOLD_FONT)
    H3   = _s(name="ZH3",   parent=base["Heading3"], fontSize=13, leading=17, spaceBefore=12, spaceAfter=3,  textColor=INK,   fontName=BOLD_FONT)
    H4   = _s(name="ZH4",   parent=base["Heading4"], fontSize=11, leading=15, spaceBefore=10, spaceAfter=2,  textColor=ACCENT, fontName=BOLD_FONT)
    BP   = _s(name="ZBP",   fontSize=11, leading=17, spaceAfter=7,  textColor=BODY, fontName=BODY_FONT)
    LI   = _s(name="ZLI",   fontSize=11, leading=17, spaceAfter=3,  textColor=BODY, fontName=BODY_FONT, leftIndent=16, bulletIndent=4)
    SIG  = _s(name="ZSig",  fontSize=10, leading=14, spaceAfter=2,  textColor=INK,  fontName=BOLD_FONT)
    SIG2 = _s(name="ZSig2", fontSize=9,  leading=13, spaceAfter=1,  textColor=MUTED, fontName=BODY_FONT)
    HDR  = _s(name="ZHdr",  fontSize=8,  leading=11, spaceAfter=0,  textColor=MUTED, fontName=BODY_FONT, alignment=2)
    CODE = _s(name="ZCode", fontSize=9,  leading=13, spaceAfter=8,  textColor=colors.HexColor("#1E293B"),
              fontName="Courier", leftIndent=12, backColor=CODE_BG, borderPadding=6)
    FTR  = _s(name="ZFtr",  fontSize=8.5, leading=12, textColor=MUTED, fontName=BODY_FONT, alignment=1)

    story = []

    # Letterhead: logo + business name band (ReportLab fallback path)
    raw_logo = (style.get("logo_url") or style.get("default_logo_url") or "").strip()
    if raw_logo or business_name:
        if raw_logo:
            embedded = _embed_image_as_data_uri(raw_logo)
            if embedded.startswith("data:"):
                import base64
                import io
                from reportlab.platypus import Image as RLImage

                m = re.match(r"data:image/[^;]+;base64,(.+)", embedded)
                if m:
                    try:
                        story.append(RLImage(io.BytesIO(base64.b64decode(m.group(1))), width=4.2 * cm, height=1.5 * cm))
                        story.append(Spacer(1, 6))
                    except Exception:
                        pass
        if business_name:
            bn_style = _s(name="ZBrand", fontName=BOLD_FONT, fontSize=14, textColor=ACCENT, spaceAfter=2)
            story.append(Paragraph(business_name, bn_style))
        story.append(HRFlowable(width="100%", thickness=2, color=ACCENT, spaceAfter=14))

    # ── Header text (top-right, from style profile) ──────────────────────────
    header_text = (style.get("header_text") or "").strip()
    if header_text:
        story.append(Paragraph(header_text, HDR))
        story.append(Spacer(1, 6))

    html = _md_to_html(markdown_content)
    blocks = _parse_blocks(html)

    # inline bold/italic helper
    def _inline(text: str) -> str:
        text = re.sub(r"\*\*(.*?)\*\*", r"<b>\1</b>", text)
        text = re.sub(r"\*(.*?)\*", r"<i>\1</i>", text)
        text = re.sub(r"`(.*?)`", r"<font name='Courier' size='9'>\1</font>", text)
        return text

    for kind, raw in blocks:
        if kind == "h1":
            story.append(Paragraph(_strip(raw), H1))
            story.append(HRFlowable(width="100%", thickness=1.2, color=ACCENT, spaceAfter=10))
        elif kind == "h2":
            story.append(Paragraph(_strip(raw), H2))
            story.append(HRFlowable(width="100%", thickness=0.5, color=RULE, spaceAfter=6))
        elif kind == "h3":
            story.append(Paragraph(_strip(raw), H3))
        elif kind == "h4":
            story.append(Paragraph(_strip(raw), H4))
        elif kind == "hr":
            story.append(Spacer(1, 4))
            story.append(HRFlowable(width="100%", thickness=0.5, color=RULE, spaceAfter=8))
        elif kind == "code":
            story.append(Spacer(1, 2))
            story.append(Paragraph(raw.replace("\n", "<br/>"), CODE))
        elif kind == "ul":
            story.append(Paragraph(f"• {_inline(raw)}", LI))
        elif kind == "ol":
            story.append(Paragraph(_inline(raw), LI))
        elif kind == "table":
            rows = _parse_table(raw)
            if rows:
                ncols = max(len(r) for r in rows)
                rows = [r + [""] * (ncols - len(r)) for r in rows]
                avail_w = A4[0] - 4.4 * cm
                col_w = avail_w / ncols
                # Wrap cells in Paragraph so long text flows instead of overflowing
                hdr_style = ParagraphStyle("TblH", fontName=BOLD_FONT, fontSize=10,
                                           leading=14, textColor=INK)
                cell_style = ParagraphStyle("TblC", fontName=BODY_FONT, fontSize=10,
                                            leading=14, textColor=INK)
                para_rows = []
                for ri, row in enumerate(rows):
                    st = hdr_style if ri == 0 else cell_style
                    para_rows.append([Paragraph(str(cell), st) for cell in row])
                tbl = Table(para_rows, colWidths=[col_w] * ncols, repeatRows=1,
                            hAlign="LEFT")
                tbl.setStyle(TableStyle([
                    ("BACKGROUND",    (0, 0), (-1, 0),  TABLE_HEAD),
                    ("LINEBELOW",     (0, 0), (-1, 0),  1.5, ACCENT),
                    ("ROWBACKGROUNDS",(0, 1), (-1, -1), [colors.white, TABLE_ALT]),
                    ("GRID",          (0, 0), (-1, -1), 0.4, RULE),
                    ("VALIGN",        (0, 0), (-1, -1), "TOP"),
                    ("TOPPADDING",    (0, 0), (-1, -1), 8),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
                    ("LEFTPADDING",   (0, 0), (-1, -1), 10),
                    ("RIGHTPADDING",  (0, 0), (-1, -1), 10),
                ]))
                story.append(tbl)
                story.append(Spacer(1, 10))
        elif kind == "p":
            story.append(Paragraph(_inline(raw), BP))

    # ── Signature block (from style profile) ────────────────────────────────
    sig_name    = (style.get("signature_name") or "").strip()
    sig_title   = (style.get("signature_title") or "").strip()
    sig_contact = (style.get("signature_contact") or "").strip()
    if sig_name or sig_title or sig_contact:
        story.append(Spacer(1, 20))
        story.append(HRFlowable(width="40%", thickness=0.6, color=ACCENT, spaceAfter=8))
        if sig_name:
            story.append(Paragraph(sig_name, SIG))
        if sig_title:
            story.append(Paragraph(sig_title, SIG2))
        if sig_contact:
            story.append(Paragraph(sig_contact, SIG2))

    # ── Footer ───────────────────────────────────────────────────────────────
    story.append(Spacer(1, 20))
    story.append(HRFlowable(width="100%", thickness=0.4, color=RULE, spaceAfter=6))
    footer_text = (style.get("footer_text") or "").strip()
    footer_line = footer_text or f"Generated by <b>Zilo Chat</b> · {business_name or ''}"
    footer_line += f" · {datetime.utcnow().strftime('%d %b %Y')}"
    story.append(Paragraph(footer_line, FTR))

    doc.build(story)
    return str(filepath)


# ─────────────────────────────────────────────────────────────────────────────
# HTML document template (Jinja2) — used for preview + WeasyPrint PDF
# ─────────────────────────────────────────────────────────────────────────────

_HTML_TEMPLATE = """<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{{ title }}</title>
<link rel="preconnect" href="https://fonts.googleapis.com">
<link rel="stylesheet" href="https://fonts.googleapis.com/css2?family={{ google_font }}:wght@400;600;700&display=swap">
<style>
  :root {
    --primary:   {{ primary_color }};
    --secondary: {{ secondary_color }};
    --ink:       #111827;
    --body:      #374151;
    --muted:     #6B7280;
    --rule:      #E5E7EB;
    --code-bg:   #F3F4F6;
    --table-alt: #F9FAFB;
  }

  * { box-sizing: border-box; margin: 0; padding: 0; }

  body {
    font-family: '{{ google_font }}', {{ font_stack }};
    font-size: 11pt;
    line-height: 1.7;
    color: var(--body);
    background: #fff;
    padding: 0;
  }

  .page {
    max-width: 760px;
    margin: 0 auto;
    padding: 52px 60px;
    overflow-x: hidden;
  }

  /* Classic letterhead (loan letters, formal business docs) */
  .classic-letterhead { margin-bottom: 32px; }
  .classic-logo {
    display: block;
    height: 52px;
    max-width: 180px;
    object-fit: contain;
    margin: 0 auto 14px;
  }
  .classic-brand {
    font-size: 15pt;
    font-weight: 700;
    color: var(--primary);
    text-align: left;
  }
  .classic-rule {
    border: none;
    border-top: 3px solid var(--primary);
    margin: 10px 0 12px;
  }
  .classic-meta-row {
    display: flex;
    justify-content: flex-end;
  }
  .classic-meta-right {
    text-align: right;
    font-size: 9pt;
    color: var(--muted);
    line-height: 1.6;
  }
  .classic-tagline { color: var(--muted); margin-bottom: 2px; }
  .classic-title {
    display: block;
    text-align: center;
    font-size: 22pt;
    font-weight: 700;
    color: var(--ink);
    margin: 4px 0 0;
    line-height: 1.25;
  }

  /* Header bar */
  .doc-header {
    display: flex;
    justify-content: space-between;
    align-items: flex-start;
    padding-bottom: 20px;
    border-bottom: 3px solid var(--primary);
    margin-bottom: 36px;
  }
  .doc-header-left h1 {
    font-size: 22pt;
    font-weight: 700;
    color: var(--ink);
    line-height: 1.2;
    margin-bottom: 4px;
  }
  .doc-header-left h1.doc-title {
    display: block;
  }
  .doc-header-left .business-name {
    font-size: 10pt;
    color: var(--muted);
    font-weight: 400;
  }
  .doc-header-right {
    text-align: right;
    font-size: 9pt;
    color: var(--muted);
    line-height: 1.8;
  }
  .doc-header-right .header-text {
    font-weight: 600;
    color: var(--primary);
  }

  /* Body content */
  h1 { display: none; } /* hide duplicate markdown h1 */
  h1.doc-title {
    display: block;
    font-size: 22pt;
    font-weight: 700;
    color: var(--ink);
    line-height: 1.2;
    margin-bottom: 4px;
  }

  h2 {
    font-size: 14pt;
    font-weight: 700;
    color: var(--ink);
    margin-top: 32px;
    margin-bottom: 8px;
    padding-bottom: 6px;
    border-bottom: 1px solid var(--rule);
  }
  h3 {
    font-size: 12pt;
    font-weight: 600;
    color: var(--ink);
    margin-top: 22px;
    margin-bottom: 6px;
  }
  h4 {
    font-size: 11pt;
    font-weight: 600;
    color: var(--primary);
    margin-top: 16px;
    margin-bottom: 4px;
  }
  p {
    margin-bottom: 10px;
    color: var(--body);
  }
  ul, ol {
    margin: 6px 0 12px 22px;
  }
  li {
    margin-bottom: 4px;
  }
  hr {
    border: none;
    border-top: 1px solid var(--rule);
    margin: 24px 0;
  }
  strong { color: var(--ink); }

  /* Tables */
  .table-wrap { width: 100%; overflow-x: auto; -webkit-overflow-scrolling: touch; }
  table {
    width: 100%;
    border-collapse: collapse;
    margin: 16px 0 20px;
    font-size: 10pt;
    table-layout: auto;
  }
  thead tr {
    background: var(--secondary);
    border-bottom: 2px solid var(--primary);
  }
  thead th {
    padding: 10px 14px;
    text-align: left;
    font-weight: 700;
    color: var(--ink);
    letter-spacing: 0.01em;
  }
  tbody tr { border-bottom: 1px solid var(--rule); }
  tbody tr:nth-child(even) { background: var(--table-alt); }
  tbody td {
    padding: 9px 14px;
    vertical-align: top;
    line-height: 1.55;
  }

  /* Code */
  pre, code {
    font-family: 'Courier New', Courier, monospace;
    font-size: 9pt;
    background: var(--code-bg);
    border-radius: 4px;
    padding: 2px 6px;
    color: #1E293B;
  }
  pre { padding: 12px 16px; margin: 12px 0; display: block; }

  /* Signature block */
  .signature {
    margin-top: 40px;
    padding-top: 16px;
    border-top: 2px solid var(--primary);
    width: 260px;
  }
  .signature .sig-name {
    font-size: 11pt;
    font-weight: 700;
    color: var(--ink);
    margin-bottom: 2px;
  }
  .signature .sig-title,
  .signature .sig-contact {
    font-size: 9.5pt;
    color: var(--muted);
    line-height: 1.6;
  }

  /* Hero image banner (professional template) */
  .doc-hero {
    width: 100%;
    height: 180px;
    object-fit: cover;
    display: block;
    margin-bottom: 0;
  }

  /* Footer */
  .doc-footer {
    margin-top: 48px;
    padding-top: 10px;
    border-top: 1px solid var(--rule);
    font-size: 8.5pt;
    color: var(--muted);
    text-align: center;
  }

  /* Print / PDF overrides */
  @media print {
    body { padding: 0; }
    .page { padding: 0; }
  }
</style>
</head>
<body>
<div class="page">

  {% if letterhead %}
  <div class="classic-letterhead">
    {% if logo_url %}<img src="{{ logo_url }}" alt="{{ business_name }}" class="classic-logo">{% endif %}
    <div class="classic-brand">{{ business_name }}</div>
    <hr class="classic-rule">
    {% if header_text or header_contact %}
    <div class="classic-meta-row">
      <div class="classic-meta-right">
        {% if header_text %}<div class="classic-tagline">{{ header_text }}</div>{% endif %}
        {% if header_contact %}<div class="classic-tagline">{{ header_contact }}</div>{% endif %}
        <div class="classic-date">{{ date_str }}</div>
      </div>
    </div>
    {% endif %}
    <h1 class="classic-title">{{ title }}</h1>
    <hr class="classic-rule">
  </div>
  {% else %}

  {% if hero_image_url %}
  <img src="{{ hero_image_url }}" class="doc-hero" alt="">
  {% endif %}

  <!-- Document header -->
  <div class="doc-header">
    <div class="doc-header-left">
      {% if logo_url %}<img src="{{ logo_url }}" alt="{{ business_name }}" style="height:36px;max-width:140px;object-fit:contain;margin-bottom:8px;display:block;">{% endif %}
      <h1 class="doc-title">{{ title }}</h1>
      <div class="business-name">{{ business_name }}</div>
    </div>
    <div class="doc-header-right">
      {% if header_text %}<div class="header-text">{{ header_text }}</div>{% endif %}
      <div>{{ date_str }}</div>
    </div>
  </div>
  {% endif %}

  <!-- Document body -->
  <div class="doc-body">
    {{ body_html }}
  </div>

  <!-- Signature -->
  {% if signature_name or signature_title or signature_contact %}
  <div class="signature">
    {% if signature_name %}<div class="sig-name">{{ signature_name }}</div>{% endif %}
    {% if signature_title %}<div class="sig-title">{{ signature_title }}</div>{% endif %}
    {% if signature_contact %}<div class="sig-contact">{{ signature_contact }}</div>{% endif %}
  </div>
  {% endif %}

  <!-- Footer -->
  <div class="doc-footer">
    {% if footer_text %}{{ footer_text }}{% else %}Generated by <strong>Zilo Chat</strong> · {{ business_name }}{% endif %}
    &nbsp;·&nbsp; {{ date_str }}
  </div>

</div>
</body>
</html>"""


_HTML_TEMPLATE_MINIMAL = """<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{{ title }}</title>
<link rel="preconnect" href="https://fonts.googleapis.com">
<link rel="stylesheet" href="https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600&display=swap">
<style>
  :root {
    --accent: {{ primary_color }};
    --ink:    #0F172A;
    --body:   #334155;
    --muted:  #94A3B8;
    --rule:   #E2E8F0;
  }
  * { box-sizing: border-box; margin: 0; padding: 0; }
  body {
    font-family: 'Inter', system-ui, sans-serif;
    font-size: 10.5pt;
    line-height: 1.75;
    color: var(--body);
    background: #F8FAFC;
  }
  .page {
    max-width: 760px;
    margin: 0 auto;
    background: #fff;
    padding: 56px 68px;
    overflow-x: hidden;
  }
  .doc-header {
    display: flex;
    justify-content: space-between;
    align-items: flex-end;
    margin-bottom: 40px;
    padding-bottom: 16px;
    border-bottom: 1px solid var(--rule);
  }
  .doc-header-left .doc-title {
    font-size: 20pt;
    font-weight: 600;
    color: var(--ink);
    letter-spacing: -0.02em;
  }
  .doc-header-left .business-name {
    font-size: 9pt;
    color: var(--muted);
    margin-top: 2px;
  }
  .doc-header-right {
    text-align: right;
    font-size: 8.5pt;
    color: var(--muted);
    line-height: 1.7;
  }
  h1 { display: none; }
  h2 {
    font-size: 9.5pt;
    font-weight: 600;
    color: var(--ink);
    text-transform: uppercase;
    letter-spacing: 0.07em;
    margin-top: 30px;
    margin-bottom: 8px;
  }
  h3 {
    font-size: 10.5pt;
    font-weight: 600;
    color: var(--ink);
    margin-top: 20px;
    margin-bottom: 5px;
  }
  h4 {
    font-size: 10pt;
    font-weight: 500;
    color: var(--accent);
    margin-top: 14px;
    margin-bottom: 3px;
  }
  p { margin-bottom: 10px; }
  ul, ol { margin: 5px 0 10px 20px; }
  li { margin-bottom: 3px; }
  hr { border: none; border-top: 1px solid var(--rule); margin: 20px 0; }
  strong { color: var(--ink); font-weight: 600; }
  .table-wrap { width: 100%; overflow-x: auto; -webkit-overflow-scrolling: touch; }
  table {
    width: 100%;
    border-collapse: collapse;
    margin: 14px 0 18px;
    font-size: 9.5pt;
    table-layout: auto;
  }
  thead th {
    padding: 9px 12px;
    text-align: left;
    font-weight: 600;
    color: var(--ink);
    border-bottom: 1.5px solid var(--ink);
  }
  tbody tr { border-bottom: 1px solid var(--rule); }
  tbody td {
    padding: 8px 12px;
    vertical-align: top;
  }
  pre, code {
    font-family: 'Courier New', monospace;
    font-size: 8.5pt;
    background: #F1F5F9;
    border-radius: 3px;
    padding: 2px 5px;
  }
  pre { padding: 10px 14px; margin: 10px 0; display: block; }
  .signature {
    margin-top: 36px;
    padding-top: 14px;
    border-top: 1px solid var(--rule);
    width: 240px;
  }
  .sig-name { font-size: 10pt; font-weight: 600; color: var(--ink); }
  .sig-title, .sig-contact { font-size: 9pt; color: var(--muted); }
  .doc-footer {
    margin-top: 44px;
    padding-top: 10px;
    border-top: 1px solid var(--rule);
    font-size: 8pt;
    color: var(--muted);
    text-align: center;
  }
  @media print { body { background: #fff; } .page { padding: 0; } }
</style>
</head>
<body>
<div class="page">
  <div class="doc-header">
    <div class="doc-header-left">
      {% if logo_url %}<img src="{{ logo_url }}" alt="{{ business_name }}" style="height:28px;max-width:120px;object-fit:contain;margin-bottom:6px;display:block;">{% endif %}
      <div class="doc-title">{{ title }}</div>
      <div class="business-name">{{ business_name }}</div>
    </div>
    <div class="doc-header-right">
      {% if header_text %}<div>{{ header_text }}</div>{% endif %}
      <div>{{ date_str }}</div>
    </div>
  </div>
  <div class="doc-body">{{ body_html }}</div>
  {% if signature_name or signature_title or signature_contact %}
  <div class="signature">
    {% if signature_name %}<div class="sig-name">{{ signature_name }}</div>{% endif %}
    {% if signature_title %}<div class="sig-title">{{ signature_title }}</div>{% endif %}
    {% if signature_contact %}<div class="sig-contact">{{ signature_contact }}</div>{% endif %}
  </div>
  {% endif %}
  <div class="doc-footer">
    {% if footer_text %}{{ footer_text }}{% else %}Prepared by <strong>{{ business_name }}</strong> · Powered by Zilo Chat{% endif %}
    &nbsp;·&nbsp; {{ date_str }}
  </div>
</div>
</body>
</html>"""


_HTML_TEMPLATE_EXECUTIVE = """<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{{ title }}</title>
<link rel="preconnect" href="https://fonts.googleapis.com">
<link rel="stylesheet" href="https://fonts.googleapis.com/css2?family=Playfair+Display:wght@600;700&family=Source+Sans+3:wght@400;600&display=swap">
<style>
  :root {
    --primary:  {{ primary_color }};
    --ink:      #0F172A;
    --body:     #1E293B;
    --muted:    #64748B;
    --rule:     #CBD5E1;
    --dark-hdr: #0F172A;
    --table-alt:#F8FAFC;
  }
  * { box-sizing: border-box; margin: 0; padding: 0; }
  body {
    font-family: 'Source Sans 3', 'Helvetica Neue', sans-serif;
    font-size: 11pt;
    line-height: 1.7;
    color: var(--body);
    background: #fff;
  }
  .page {
    max-width: 760px;
    margin: 0 auto;
    padding: 0 0 52px 0;
    overflow-x: hidden;
  }
  .doc-header {
    position: relative;
    background: var(--dark-hdr);
    color: #fff;
    padding: 36px 52px 28px;
    margin-bottom: 36px;
    overflow: hidden;
  }
  .doc-header-bg {
    position: absolute;
    inset: 0;
    background-size: cover;
    background-position: center;
    {% if hero_image_url %}background-image: url('{{ hero_image_url }}');{% endif %}
    opacity: 0.22;
  }
  .doc-header-overlay {
    position: absolute;
    inset: 0;
    background: linear-gradient(135deg, rgba(15,23,42,0.97) 0%, rgba(15,23,42,0.80) 100%);
  }
  .doc-header-content {
    position: relative;
    z-index: 2;
  }
  .doc-header-top {
    display: flex;
    justify-content: space-between;
    align-items: flex-start;
  }
  .doc-header-left .doc-title {
    font-family: 'Playfair Display', Georgia, serif;
    font-size: 24pt;
    font-weight: 700;
    color: #fff;
    line-height: 1.2;
    margin-bottom: 6px;
  }
  .doc-header-left .business-name {
    font-size: 10pt;
    color: rgba(255,255,255,0.5);
  }
  .doc-header-right {
    text-align: right;
    font-size: 9pt;
    color: rgba(255,255,255,0.5);
    line-height: 1.8;
  }
  .doc-header-right .header-text {
    font-weight: 600;
    color: var(--primary);
  }
  .doc-header-divider {
    border: none;
    border-top: 1px solid rgba(255,255,255,0.12);
    margin-top: 20px;
  }
  .doc-body { padding: 0 52px; }
  h1 { display: none; }
  h2 {
    font-family: 'Playfair Display', Georgia, serif;
    font-size: 15pt;
    font-weight: 600;
    color: var(--ink);
    margin-top: 32px;
    margin-bottom: 10px;
    padding-bottom: 6px;
    border-bottom: 2px solid var(--primary);
  }
  h3 {
    font-size: 12pt;
    font-weight: 600;
    color: var(--ink);
    margin-top: 22px;
    margin-bottom: 6px;
  }
  h4 {
    font-size: 10.5pt;
    font-weight: 600;
    color: var(--primary);
    margin-top: 15px;
    margin-bottom: 4px;
  }
  p { margin-bottom: 10px; }
  ul, ol { margin: 6px 0 12px 22px; }
  li { margin-bottom: 4px; }
  hr { border: none; border-top: 1px solid var(--rule); margin: 22px 0; }
  strong { color: var(--ink); font-weight: 600; }
  .table-wrap { width: 100%; overflow-x: auto; -webkit-overflow-scrolling: touch; }
  table {
    width: 100%;
    border-collapse: collapse;
    margin: 16px 0 20px;
    font-size: 10pt;
    table-layout: auto;
  }
  thead tr { background: var(--dark-hdr); }
  thead th {
    padding: 10px 14px;
    text-align: left;
    font-weight: 600;
    color: #fff;
    letter-spacing: 0.02em;
  }
  tbody tr { border-bottom: 1px solid var(--rule); }
  tbody tr:nth-child(even) { background: var(--table-alt); }
  tbody td {
    padding: 9px 14px;
    vertical-align: top;
  }
  pre, code {
    font-family: 'Courier New', monospace;
    font-size: 9pt;
    background: #F1F5F9;
    border-radius: 4px;
    padding: 2px 6px;
  }
  pre { padding: 12px 16px; margin: 12px 0; display: block; }
  .signature {
    margin: 40px 52px 0;
    padding-top: 16px;
    border-top: 2px solid var(--primary);
    width: 260px;
  }
  .sig-name { font-size: 11pt; font-weight: 600; color: var(--ink); margin-bottom: 2px; }
  .sig-title, .sig-contact { font-size: 9.5pt; color: var(--muted); }
  .doc-footer {
    margin: 44px 52px 0;
    padding-top: 10px;
    border-top: 1px solid var(--rule);
    font-size: 8.5pt;
    color: var(--muted);
    text-align: center;
  }
  @media print { .page { padding-bottom: 0; } }
</style>
</head>
<body>
<div class="page">
  <div class="doc-header">
    <div class="doc-header-bg"></div>
    <div class="doc-header-overlay"></div>
    <div class="doc-header-content">
      <div class="doc-header-top">
        <div class="doc-header-left">
          {% if logo_url %}<img src="{{ logo_url }}" alt="{{ business_name }}" style="height:32px;max-width:130px;object-fit:contain;margin-bottom:10px;display:block;filter:brightness(0) invert(1);opacity:0.9;">{% endif %}
          <div class="doc-title">{{ title }}</div>
          <div class="business-name">{{ business_name }}</div>
        </div>
        <div class="doc-header-right">
          {% if header_text %}<div class="header-text">{{ header_text }}</div>{% endif %}
          <div>{{ date_str }}</div>
        </div>
      </div>
      <hr class="doc-header-divider">
    </div>
  </div>
  <div class="doc-body">{{ body_html }}</div>
  {% if signature_name or signature_title or signature_contact %}
  <div class="signature">
    {% if signature_name %}<div class="sig-name">{{ signature_name }}</div>{% endif %}
    {% if signature_title %}<div class="sig-title">{{ signature_title }}</div>{% endif %}
    {% if signature_contact %}<div class="sig-contact">{{ signature_contact }}</div>{% endif %}
  </div>
  {% endif %}
  <div class="doc-footer">
    {% if footer_text %}{{ footer_text }}{% else %}Generated by <strong>Zilo Chat</strong> · {{ business_name }}{% endif %}
    &nbsp;·&nbsp; {{ date_str }}
  </div>
</div>
</body>
</html>"""


_TEMPLATE_MAP = {
    "professional": _HTML_TEMPLATE,
    "minimal":      _HTML_TEMPLATE_MINIMAL,
    "executive":    _HTML_TEMPLATE_EXECUTIVE,
}


def _font_stack_for_style(font_style: str) -> tuple[str, str]:
    """Return (google_font_name, css_generic_stack)."""
    s = (font_style or "").lower()
    if s in ("serif", "classic"):
        return "Libre Baskerville", "Georgia, 'Times New Roman', serif"
    if s == "modern":
        return "DM Sans", "'Helvetica Neue', Arial, sans-serif"
    if s == "minimal":
        return "Inter", "system-ui, sans-serif"
    # default sans-serif / professional
    return "Inter", "'Helvetica Neue', Arial, sans-serif"


def generate_html_document(
    markdown_content: str,
    title: str = "Document",
    business_name: str = "",
    style: Optional[Dict[str, Any]] = None,
    template: str = "professional",
    hero_image_url: Optional[str] = None,
) -> str:
    """Render the document as a self-contained HTML string for preview or WeasyPrint."""
    from jinja2 import Template

    style = style or {}
    primary_color   = _safe_hex(style.get("primary_color", ""), "#4F46E5")
    secondary_color = _safe_hex(style.get("secondary_color", ""), "#EEF2FF")
    google_font, font_stack = _font_stack_for_style(style.get("font_style", ""))

    # Logo from brand kit — embed as data URI so preview iframe + PDF render reliably
    raw_logo: str = style.get("logo_url", "") or style.get("default_logo_url", "") or ""
    logo_url: str = _embed_image_as_data_uri(raw_logo) if raw_logo else ""

    if hero_image_url and not hero_image_url.startswith("data:"):
        hero_image_url = _embed_image_as_data_uri(hero_image_url) or hero_image_url

    # Convert markdown body (strip leading h1 — title shown in header)
    body_md = re.sub(r"^\s*#[^#][^\n]*\n?", "", markdown_content, count=1).strip()
    body_html = _md_to_html(body_md)
    # Strip standalone <img> paragraphs at the start of the body — these are logo
    # images the AI occasionally places in the markdown. They belong in the header,
    # not floating in the middle of the document.
    body_html = re.sub(r"^\s*<p>\s*<img[^>]+>\s*</p>\s*", "", body_html, count=3)
    # Wrap every <table> in a .table-wrap div so long rows never overflow the page
    body_html = re.sub(r"(<table>)", r'<div class="table-wrap">\1', body_html)
    body_html = re.sub(r"(</table>)", r"\1</div>", body_html)

    date_fmt = style.get("date_format", "") or "DD Month YYYY"
    now = datetime.utcnow()
    if "YYYY" in date_fmt and "Month" in date_fmt:
        date_str = now.strftime("%-d %B %Y") if os.name != "nt" else now.strftime("%d %B %Y").lstrip("0")
    elif "MM/DD/YYYY" in date_fmt:
        date_str = now.strftime("%m/%d/%Y")
    elif "DD/MM/YYYY" in date_fmt:
        date_str = now.strftime("%d/%m/%Y")
    else:
        date_str = now.strftime("%d %B %Y")

    tmpl_str = _TEMPLATE_MAP.get(template, _HTML_TEMPLATE)
    tmpl = Template(tmpl_str)
    letterhead = bool(style.get("letterhead"))
    return tmpl.render(
        title=title,
        business_name=business_name or "My Business",
        primary_color=primary_color,
        secondary_color=secondary_color,
        google_font=google_font,
        font_stack=font_stack,
        header_text=style.get("header_text", ""),
        header_contact=style.get("header_contact", ""),
        footer_text=style.get("footer_text", ""),
        signature_name=style.get("signature_name", ""),
        signature_title=style.get("signature_title", ""),
        signature_contact=style.get("signature_contact", ""),
        date_str=date_str,
        body_html=body_html,
        hero_image_url=hero_image_url or "",
        logo_url=logo_url,
        letterhead=letterhead,
    )


def _pdf_exc_detail(exc: BaseException) -> str:
    """Readable one-line error for logs and user-facing messages."""
    msg = str(exc).strip()
    if msg:
        first = msg.splitlines()[0].strip()
        return first if len(first) <= 240 else first[:237] + "..."
    return f"{type(exc).__name__} (no message)"


def _pydyf_version() -> str:
    try:
        import pydyf
        return getattr(pydyf, "__version__", "") or ""
    except ImportError:
        return ""


def _pydyf_is_compatible() -> bool:
    ver = _pydyf_version()
    if not ver:
        return True
    parts = ver.split(".")
    try:
        major, minor = int(parts[0]), int(parts[1]) if len(parts) > 1 else 0
    except ValueError:
        return True
    return (major, minor) < (0, 11)


def _resolve_pdf_path(filename: Optional[str]) -> Path:
    name = filename or f"zilo_{uuid.uuid4().hex[:8]}.pdf"
    if not name.endswith(".pdf"):
        name += ".pdf"
    return TEMP_DIR / name


def _playwright_pdf_sync(html: str, filepath: Path) -> None:
    """Sync Playwright render — safe inside a subprocess or worker thread."""
    from playwright.sync_api import sync_playwright

    with sync_playwright() as p:
        browser = p.chromium.launch(
            headless=True,
            args=["--no-sandbox", "--disable-dev-shm-usage"],
        )
        try:
            page = browser.new_page(viewport={"width": 1200, "height": 900})
            try:
                page.set_content(html, wait_until="domcontentloaded", timeout=60_000)
            except Exception:
                page.set_content(html, wait_until="commit", timeout=30_000)
            page.emulate_media(media="print")
            page.pdf(
                path=str(filepath),
                format="A4",
                print_background=True,
                margin={"top": "0mm", "right": "0mm", "bottom": "0mm", "left": "0mm"},
            )
        finally:
            browser.close()


def _playwright_pdf_subprocess(html: str, filepath: Path) -> None:
    """Run Playwright in a fresh Python process (avoids asyncio/event-loop conflicts)."""
    script = r"""
import sys
from pathlib import Path
from playwright.sync_api import sync_playwright

html = sys.stdin.read()
out = Path(sys.argv[1])
with sync_playwright() as p:
    browser = p.chromium.launch(headless=True, args=["--no-sandbox", "--disable-dev-shm-usage"])
    try:
        page = browser.new_page(viewport={"width": 1200, "height": 900})
        try:
            page.set_content(html, wait_until="domcontentloaded", timeout=60000)
        except Exception:
            page.set_content(html, wait_until="commit", timeout=30000)
        page.emulate_media(media="print")
        page.pdf(
            path=str(out),
            format="A4",
            print_background=True,
            margin={"top": "0mm", "right": "0mm", "bottom": "0mm", "left": "0mm"},
        )
    finally:
        browser.close()
"""
    proc = subprocess.run(
        [sys.executable, "-c", script, str(filepath)],
        input=html.encode("utf-8"),
        capture_output=True,
        timeout=120,
        cwd=str(ROOT_DIR),
    )
    if proc.returncode != 0:
        err = (proc.stderr or proc.stdout or b"").decode("utf-8", errors="replace").strip()
        first = err.splitlines()[0].strip() if err else f"exit code {proc.returncode}"
        raise RuntimeError(first or "Playwright subprocess failed")


def _weasyprint_pdf(html: str, filepath: Path) -> None:
    if not _pydyf_is_compatible():
        ver = _pydyf_version()
        raise RuntimeError(
            f"Incompatible pydyf {ver} for weasyprint 62.x — run: pip install pydyf==0.10.0"
        )
    import weasyprint

    weasyprint.HTML(string=html).write_pdf(str(filepath))


async def generate_pdf_from_html_async(
    html: str,
    filename: Optional[str] = None,
) -> str:
    """Convert HTML to PDF — Playwright (async → subprocess → sync thread), then WeasyPrint."""
    import asyncio
    import logging

    filepath = _resolve_pdf_path(filename)
    playwright_err = ""

    # 1) Playwright async API (fast path inside uvicorn)
    try:
        from playwright.async_api import async_playwright

        async with async_playwright() as p:
            browser = await p.chromium.launch(
                headless=True,
                args=["--no-sandbox", "--disable-dev-shm-usage"],
            )
            try:
                page = await browser.new_page(viewport={"width": 1200, "height": 900})
                try:
                    await page.set_content(html, wait_until="domcontentloaded", timeout=60_000)
                except Exception:
                    await page.set_content(html, wait_until="commit", timeout=30_000)
                await page.emulate_media(media="print")
                await page.pdf(
                    path=str(filepath),
                    format="A4",
                    print_background=True,
                    margin={"top": "0mm", "right": "0mm", "bottom": "0mm", "left": "0mm"},
                )
            finally:
                await browser.close()
        return str(filepath)
    except Exception as exc:
        playwright_err = _pdf_exc_detail(exc)
        logging.warning("[document_generator] Playwright async PDF failed: %s", playwright_err)

    # 2) Isolated subprocess — avoids asyncio / Windows event-loop issues
    try:
        await asyncio.to_thread(_playwright_pdf_subprocess, html, filepath)
        logging.info("[document_generator] PDF generated via Playwright subprocess")
        return str(filepath)
    except Exception as exc:
        sub_err = _pdf_exc_detail(exc)
        logging.warning("[document_generator] Playwright subprocess PDF failed: %s", sub_err)
        if not playwright_err:
            playwright_err = sub_err

    # 3) Sync Playwright in a dedicated thread (no running asyncio loop)
    try:
        await asyncio.to_thread(_playwright_pdf_sync, html, filepath)
        logging.info("[document_generator] PDF generated via Playwright sync thread")
        return str(filepath)
    except Exception as exc:
        sync_err = _pdf_exc_detail(exc)
        logging.warning("[document_generator] Playwright sync PDF failed: %s", sync_err)
        if not playwright_err:
            playwright_err = sync_err

    # 4) WeasyPrint fallback
    try:
        await asyncio.to_thread(_weasyprint_pdf, html, filepath)
        logging.info("[document_generator] PDF generated via WeasyPrint fallback")
        return str(filepath)
    except ImportError:
        pass
    except Exception as exc:
        logging.warning("[document_generator] WeasyPrint PDF failed: %s", exc)
        weasy_err = _pdf_exc_detail(exc)
        hint = "pip install pydyf==0.10.0 playwright weasyprint && playwright install chromium"
        raise RuntimeError(
            f"PDF generation failed (Playwright: {playwright_err}; WeasyPrint: {weasy_err}). "
            f"Fix: {hint}"
        ) from exc

    hint = (
        "Install PDF dependencies: pip install playwright weasyprint pydyf==0.10.0 "
        "&& playwright install chromium"
    )
    raise RuntimeError(f"PDF generation failed (Playwright: {playwright_err}). {hint}")


def generate_pdf_from_html(
    html: str,
    filename: Optional[str] = None,
) -> str:
    """Sync wrapper — delegates to async Playwright in a fresh event loop."""
    import asyncio
    import concurrent.futures

    try:
        asyncio.get_running_loop()
    except RuntimeError:
        return asyncio.run(generate_pdf_from_html_async(html, filename))

    with concurrent.futures.ThreadPoolExecutor(max_workers=1) as pool:
        return pool.submit(
            lambda: asyncio.run(generate_pdf_from_html_async(html, filename))
        ).result(timeout=120)


# ─────────────────────────────────────────────────────────────────────────────
# DOCX  (python-docx)
# ─────────────────────────────────────────────────────────────────────────────
def generate_docx(
    markdown_content: str,
    filename: str | None = None,
    business_name: str = "",
    style: dict | None = None,
) -> str:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor

    style = style or {}

    filename = filename or f"zilo_{uuid.uuid4().hex[:8]}.docx"
    if not filename.endswith(".docx"):
        filename += ".docx"
    filepath = TEMP_DIR / filename

    doc = Document()

    # ── colour palette — use saved brand colors when available ────────────
    PRIMARY_HEX = _safe_hex(style.get("primary_color", ""), "#4F46E5")
    SECONDARY_HEX = _safe_hex(style.get("secondary_color", ""), "#EEF2FF")
    pr = int(PRIMARY_HEX[1:3], 16)
    pg = int(PRIMARY_HEX[3:5], 16)
    pb = int(PRIMARY_HEX[5:7], 16)
    ACCENT = RGBColor(pr, pg, pb)

    # ── fonts — serif vs sans-serif based on style profile ────────────────
    BODY_FONT, BOLD_FONT = _resolve_fonts(style.get("font_style", ""))

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.4)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── helper: set paragraph spacing ──────────────────────────────────────
    def _spacing(para, before: int = 0, after: int = 6, line: int = 276):
        pPr = para._p.get_or_add_pPr()
        spg = OxmlElement("w:spacing")
        spg.set(qn("w:before"), str(before))
        spg.set(qn("w:after"), str(after))
        spg.set(qn("w:line"), str(line))
        spg.set(qn("w:lineRule"), "auto")
        pPr.append(spg)

    # ── helper: add horizontal rule ─────────────────────────────────────────
    def _add_rule(para, color: str = "E5E7EB"):
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), color)
        pBdr.append(bottom)
        pPr.append(pBdr)

    # ── helper: table shading ───────────────────────────────────────────────
    def _shade_cell(cell, hex_color: str):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)

    # ── Header text (from style profile) ──────────────────────────────────
    header_text = (style.get("header_text") or "").strip()
    if header_text:
        hp = doc.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = hp.add_run(header_text)
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
        run.font.name = BODY_FONT
        _spacing(hp, before=0, after=60)

    html = _md_to_html(markdown_content)
    blocks = _parse_blocks(html)

    for kind, raw in blocks:
        if kind == "h1":
            p = doc.add_heading(_strip(raw), level=1)
            p.runs[0].font.color.rgb = RGBColor(0x11, 0x18, 0x27)
            p.runs[0].font.size = Pt(22)
            p.runs[0].font.name = BOLD_FONT
            _add_rule(p, PRIMARY_HEX.lstrip("#"))
            _spacing(p, before=0, after=120)
        elif kind == "h2":
            p = doc.add_heading(_strip(raw), level=2)
            p.runs[0].font.color.rgb = RGBColor(0x11, 0x18, 0x27)
            p.runs[0].font.size = Pt(16)
            p.runs[0].font.name = BOLD_FONT
            _add_rule(p)
            _spacing(p, before=240, after=80)
        elif kind == "h3":
            p = doc.add_heading(_strip(raw), level=3)
            p.runs[0].font.size = Pt(13)
            p.runs[0].font.name = BOLD_FONT
            _spacing(p, before=160, after=60)
        elif kind == "h4":
            p = doc.add_heading(_strip(raw), level=4)
            p.runs[0].font.size = Pt(11)
            p.runs[0].font.color.rgb = ACCENT
            p.runs[0].font.name = BOLD_FONT
        elif kind == "hr":
            p = doc.add_paragraph()
            _add_rule(p)
        elif kind == "code":
            p = doc.add_paragraph(raw)
            p.style = doc.styles["Normal"]
            for run in p.runs:
                run.font.name = "Courier New"
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
            _shade_cell_like_para = OxmlElement("w:pPr")
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "F3F4F6")
            _shade_cell_like_para.append(shd)
            p._p.insert(0, _shade_cell_like_para)
        elif kind == "ul":
            p = doc.add_paragraph(raw, style="List Bullet")
            for run in p.runs:
                run.font.name = BODY_FONT
            _spacing(p, before=0, after=40)
        elif kind == "ol":
            p = doc.add_paragraph(raw, style="List Number")
            for run in p.runs:
                run.font.name = BODY_FONT
            _spacing(p, before=0, after=40)
        elif kind == "table":
            rows = _parse_table(raw)
            if rows and rows[0]:
                ncols = max(len(r) for r in rows)
                rows = [r + [""] * (ncols - len(r)) for r in rows]
                tbl = doc.add_table(rows=len(rows), cols=ncols)
                tbl.style = "Table Grid"
                for i, row_data in enumerate(rows):
                    for j, text in enumerate(row_data):
                        cell = tbl.rows[i].cells[j]
                        cell.text = text
                        para = cell.paragraphs[0]
                        if i == 0:
                            _shade_cell(cell, SECONDARY_HEX.lstrip("#"))
                            if para.runs:
                                para.runs[0].bold = True
                                para.runs[0].font.color.rgb = RGBColor(0x11, 0x18, 0x27)
                                para.runs[0].font.name = BOLD_FONT
                        elif i % 2 == 0:
                            _shade_cell(cell, "F9FAFB")
                        else:
                            for run in para.runs:
                                run.font.name = BODY_FONT
                doc.add_paragraph()
        elif kind == "p":
            p = doc.add_paragraph(raw)
            for run in p.runs:
                run.font.name = BODY_FONT
            _spacing(p, before=0, after=80)

    # ── Signature block (from style profile) ─────────────────────────────
    sig_name = (style.get("signature_name") or "").strip()
    sig_title = (style.get("signature_title") or "").strip()
    sig_contact = (style.get("signature_contact") or "").strip()
    if sig_name or sig_title or sig_contact:
        doc.add_paragraph()
        sig_rule = doc.add_paragraph()
        _add_rule(sig_rule, PRIMARY_HEX.lstrip("#"))
        _spacing(sig_rule, before=200, after=80)
        if sig_name:
            sp = doc.add_paragraph(sig_name)
            for run in sp.runs:
                run.bold = True
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0x11, 0x18, 0x27)
                run.font.name = BOLD_FONT
            _spacing(sp, before=0, after=20)
        if sig_title:
            sp = doc.add_paragraph(sig_title)
            for run in sp.runs:
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
                run.font.name = BODY_FONT
            _spacing(sp, before=0, after=10)
        if sig_contact:
            sp = doc.add_paragraph(sig_contact)
            for run in sp.runs:
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
                run.font.name = BODY_FONT
            _spacing(sp, before=0, after=10)

    # ── Footer ───────────────────────────────────────────────────────────────
    doc.add_paragraph()
    footer_rule = doc.add_paragraph()
    _add_rule(footer_rule)
    _spacing(footer_rule, before=200, after=60)
    footer_text = (style.get("footer_text") or "").strip()
    footer_line = footer_text or f"Generated by Zilo Chat · {business_name or ''}"
    footer_line += f" · {datetime.utcnow().strftime('%d %b %Y')}"
    footer_p = doc.add_paragraph(footer_line)
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if footer_p.runs:
        footer_p.runs[0].font.size = Pt(8.5)
        footer_p.runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
        footer_p.runs[0].font.name = BODY_FONT

    doc.save(str(filepath))
    return str(filepath)


# ─────────────────────────────────────────────────────────────────────────────
# Excel (.xlsx) — SaaS financial model with live formulas (openpyxl)
# ─────────────────────────────────────────────────────────────────────────────
#
# Builds a multi-sheet, formula-driven SaaS model. Every projected value is an
# Excel formula referencing the editable Assumptions sheet — change one input and
# the whole model recalculates when opened in Excel / Google Sheets / LibreOffice.
#
# Sheets: Assumptions · Revenue · P&L · CashFlow · Summary
#
# `assumptions` keys (all optional — sensible SaaS defaults applied):
#   arpa, starting_customers, new_customers_m1, new_growth, churn,
#   gross_margin, cac, opex_m1, opex_growth, starting_cash
# ─────────────────────────────────────────────────────────────────────────────

# Default monthly SaaS assumptions (USD). Tuned for an early-stage product.
_XLSX_DEFAULT_ASSUMPTIONS: dict = {
    "arpa": 49.0,              # avg revenue per account / month
    "starting_customers": 50,
    "new_customers_m1": 25,    # new customers acquired in month 1
    "new_growth": 0.08,        # month-over-month growth in new-customer adds
    "churn": 0.03,             # monthly logo churn
    "gross_margin": 0.80,      # gross margin %
    "cac": 120.0,              # customer acquisition cost
    "opex_m1": 18000.0,        # fixed opex in month 1 (salaries, tools, etc.)
    "opex_growth": 0.04,       # month-over-month opex growth
    "starting_cash": 250000.0, # cash on hand at month 0
}


def generate_xlsx(
    assumptions: dict | None = None,
    filename: str | None = None,
    business_name: str = "",
    style: dict | None = None,
    months: int = 36,
) -> str:
    """Generate a formula-driven SaaS financial model workbook. Returns the file path."""
    from openpyxl import Workbook
    from openpyxl.chart import BarChart, LineChart, Reference
    from openpyxl.formatting.rule import CellIsRule
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
    from openpyxl.utils import get_column_letter
    from openpyxl.worksheet.datavalidation import DataValidation

    a = {**_XLSX_DEFAULT_ASSUMPTIONS, **(assumptions or {})}
    style = style or {}
    months = max(1, min(int(months or 36), 120))

    filename = filename or f"zilo_model_{uuid.uuid4().hex[:8]}.xlsx"
    if not filename.endswith(".xlsx"):
        filename += ".xlsx"
    filepath = TEMP_DIR / filename

    # ── palette (brand-aware, same convention as generate_docx) ───────────────
    PRIMARY_HEX = _safe_hex(style.get("primary_color", ""), "#4F46E5").lstrip("#").upper()
    HEADER_FILL = PatternFill("solid", fgColor=PRIMARY_HEX)
    INPUT_FILL = PatternFill("solid", fgColor="FFF7E6")   # editable inputs = soft amber
    BAND_FILL = PatternFill("solid", fgColor="F3F4F6")
    WHITE = Font(color="FFFFFF", bold=True, size=11)
    BOLD = Font(bold=True)
    MUTED = Font(color="6B7280", italic=True, size=9)
    THIN = Side(style="thin", color="E5E7EB")
    BORDER = Border(left=THIN, right=THIN, top=THIN, bottom=THIN)
    CENTER = Alignment(horizontal="center", vertical="center")
    MONEY = '#,##0;[Red](#,##0)'
    MONEY2 = '$#,##0'
    PCT = '0.0%'
    INT = '#,##0'

    wb = Workbook()

    def _hdr(ws, row, headers, *, fill=HEADER_FILL, font=WHITE):
        for j, h in enumerate(headers, start=1):
            c = ws.cell(row=row, column=j, value=h)
            c.fill = fill
            c.font = font
            c.alignment = CENTER
            c.border = BORDER

    # ── Sheet 1: Assumptions (the single source of truth) ─────────────────────
    aw = wb.active
    aw.title = "Assumptions"
    aw["A1"] = f"{business_name or 'Zilo'} — SaaS Financial Model"
    aw["A1"].font = Font(bold=True, size=15, color="111827")
    aw["A2"] = "Edit the amber cells. Every other sheet recalculates from these."
    aw["A2"].font = MUTED

    # (label, key, number_format) — values land in column B, referenced as $B$<row>
    rows = [
        ("Avg revenue per account / month (ARPA)", "arpa", MONEY2),
        ("Starting customers", "starting_customers", INT),
        ("New customers — month 1", "new_customers_m1", INT),
        ("New-customer growth (MoM)", "new_growth", PCT),
        ("Monthly churn", "churn", PCT),
        ("Gross margin", "gross_margin", PCT),
        ("Customer acquisition cost (CAC)", "cac", MONEY2),
        ("Operating expenses — month 1", "opex_m1", MONEY2),
        ("Opex growth (MoM)", "opex_growth", PCT),
        ("Starting cash", "starting_cash", MONEY2),
    ]
    # Scenario tuning factors per driver: (best_multiplier, worst_multiplier).
    # "Good when higher" drivers improve >1 in Best; "good when lower" (churn, CAC,
    # opex) improve <1 in Best. Neutral drivers stay at 1.0.
    _SCN_FACTORS = {
        "arpa": (1.15, 0.90), "starting_customers": (1.0, 1.0),
        "new_customers_m1": (1.25, 0.70), "new_growth": (1.5, 0.5),
        "churn": (0.6, 1.6), "gross_margin": (1.08, 0.88),
        "cac": (0.8, 1.35), "opex_m1": (0.95, 1.15),
        "opex_growth": (0.8, 1.3), "starting_cash": (1.0, 1.0),
    }

    def _scn(key: str, mult: float) -> float:
        v = a[key] * mult
        if fmt_for[key] == PCT:  # keep rates sane
            return round(max(0.0, min(0.99, v)), 4)
        return round(v, 2) if isinstance(a[key], float) else int(round(v))

    fmt_for = {key: fmt for (_l, key, fmt) in rows}

    # Scenario selector (drives every CHOOSE below). 1=Base, 2=Best, 3=Worst.
    aw["D2"] = "Scenario →"
    aw["D2"].font = BOLD
    aw["D2"].alignment = Alignment(horizontal="right")
    sel = aw["E2"]
    sel.value = 1
    sel.fill = INPUT_FILL
    sel.border = BORDER
    sel.alignment = CENTER
    sel.font = Font(bold=True, color=PRIMARY_HEX)
    dv = DataValidation(type="list", formula1='"1,2,3"', allow_blank=False)
    aw.add_data_validation(dv)
    dv.add(sel)
    aw["F2"] = "1 = Base · 2 = Best · 3 = Worst"
    aw["F2"].font = MUTED

    # Column headers for the driver table.
    _hdr(aw, 3, ["Driver", "Active", "", "Base", "Best", "Worst"])

    key_row: dict[str, int] = {}
    r0 = 4
    for i, (label, key, fmt) in enumerate(rows):
        r = r0 + i
        key_row[key] = r
        aw.cell(row=r, column=1, value=label).font = BOLD
        # Active value = CHOOSE(selector, Base, Best, Worst) — whole model follows it.
        ac = aw.cell(row=r, column=2, value=f"=CHOOSE($E$2,D{r},E{r},F{r})")
        ac.number_format = fmt
        ac.border = BORDER
        ac.alignment = Alignment(horizontal="right")
        ac.font = BOLD
        # Editable scenario inputs (Base / Best / Worst).
        best_m, worst_m = _SCN_FACTORS[key]
        for col, val in ((4, a[key]), (5, _scn(key, best_m)), (6, _scn(key, worst_m))):
            sc = aw.cell(row=r, column=col, value=val)
            sc.number_format = fmt
            sc.fill = INPUT_FILL
            sc.border = BORDER
            sc.alignment = Alignment(horizontal="right")
    for col, w in zip("ABCDEF", (40, 14, 2, 14, 14, 14)):
        aw.column_dimensions[col].width = w

    def AS(key: str) -> str:
        """Absolute reference to the ACTIVE assumption value, e.g. Assumptions!$B$4."""
        return f"Assumptions!$B${key_row[key]}"

    # ── Sheet 2: Revenue (months as rows) ─────────────────────────────────────
    rv = wb.create_sheet("Revenue")
    _hdr(rv, 1, ["Month", "Customers (start)", "New", "Churned", "Customers (end)", "MRR", "ARR"])
    for m in range(1, months + 1):
        row = m + 1
        prev = row - 1
        rv.cell(row=row, column=1, value=m).alignment = CENTER
        if m == 1:
            rv.cell(row=row, column=2, value=f"={AS('starting_customers')}")
            rv.cell(row=row, column=3, value=f"={AS('new_customers_m1')}")
        else:
            rv.cell(row=row, column=2, value=f"=E{prev}")
            rv.cell(row=row, column=3, value=f"=ROUND(C{prev}*(1+{AS('new_growth')}),0)")
        rv.cell(row=row, column=4, value=f"=ROUND(B{row}*{AS('churn')},0)")
        rv.cell(row=row, column=5, value=f"=B{row}+C{row}-D{row}")
        rv.cell(row=row, column=6, value=f"=E{row}*{AS('arpa')}")
        rv.cell(row=row, column=7, value=f"=F{row}*12")
        for col in range(1, 8):
            cell = rv.cell(row=row, column=col)
            cell.border = BORDER
            if col in (6, 7):
                cell.number_format = MONEY2
            elif col != 1:
                cell.number_format = INT
            if m % 2 == 0:
                cell.fill = BAND_FILL
    for col, w in zip("ABCDEFG", (8, 16, 10, 10, 16, 14, 16)):
        rv.column_dimensions[col].width = w

    # ── Sheet 3: P&L (months as rows) ─────────────────────────────────────────
    pl = wb.create_sheet("P&L")
    _hdr(pl, 1, ["Month", "Revenue", "COGS", "Gross Profit", "Sales & Mktg", "Other Opex", "EBITDA"])
    for m in range(1, months + 1):
        row = m + 1
        pl.cell(row=row, column=1, value=m).alignment = CENTER
        pl.cell(row=row, column=2, value=f"=Revenue!F{row}")
        pl.cell(row=row, column=3, value=f"=B{row}*(1-{AS('gross_margin')})")
        pl.cell(row=row, column=4, value=f"=B{row}-C{row}")
        pl.cell(row=row, column=5, value=f"=Revenue!C{row}*{AS('cac')}")
        pl.cell(row=row, column=6, value=f"={AS('opex_m1')}*(1+{AS('opex_growth')})^({m}-1)")
        pl.cell(row=row, column=7, value=f"=D{row}-E{row}-F{row}")
        for col in range(1, 8):
            cell = pl.cell(row=row, column=col)
            cell.border = BORDER
            if col != 1:
                cell.number_format = MONEY
            if m % 2 == 0:
                cell.fill = BAND_FILL
    for col, w in zip("ABCDEFG", (8, 14, 14, 14, 14, 14, 14)):
        pl.column_dimensions[col].width = w

    # ── Sheet 4: Cash Flow (running balance) ──────────────────────────────────
    cf = wb.create_sheet("CashFlow")
    _hdr(cf, 1, ["Month", "Beginning Cash", "Net Cash (EBITDA)", "Ending Cash"])
    for m in range(1, months + 1):
        row = m + 1
        prev = row - 1
        cf.cell(row=row, column=1, value=m).alignment = CENTER
        cf.cell(row=row, column=2, value=(f"={AS('starting_cash')}" if m == 1 else f"=D{prev}"))
        cf.cell(row=row, column=3, value=f"='P&L'!G{row}")
        cf.cell(row=row, column=4, value=f"=B{row}+C{row}")
        for col in range(1, 5):
            cell = cf.cell(row=row, column=col)
            cell.border = BORDER
            if col != 1:
                cell.number_format = MONEY
            if m % 2 == 0:
                cell.fill = BAND_FILL
    for col, w in zip("ABCD", (8, 18, 20, 18)):
        cf.column_dimensions[col].width = w
    # Highlight any month the cash balance goes negative (runway cliff).
    red_fill = PatternFill("solid", fgColor="FECACA")
    red_font = Font(color="991B1B", bold=True)
    cf.conditional_formatting.add(
        f"D2:D{months + 1}",
        CellIsRule(operator="lessThan", formula=["0"], fill=red_fill, font=red_font),
    )

    # ── Sheet 5: Summary / dashboard ──────────────────────────────────────────
    sm = wb.create_sheet("Summary")
    sm["A1"] = f"{business_name or 'Zilo'} — Key Metrics"
    sm["A1"].font = Font(bold=True, size=15, color="111827")
    last = months + 1  # last data row across the monthly sheets
    metrics = [
        ("Projection horizon (months)", str(months), INT, False),
        ("Ending customers", f"=Revenue!E{last}", INT, False),
        ("Ending MRR", f"=Revenue!F{last}", MONEY2, False),
        ("Ending ARR", f"=Revenue!G{last}", MONEY2, True),
        ("Year-1 revenue", "=SUM('P&L'!B2:B13)", MONEY2, False),
        ("Year-1 EBITDA", "=SUM('P&L'!G2:G13)", MONEY, False),
        ("Ending cash", f"=CashFlow!D{last}", MONEY2, False),
        ("Lowest cash balance", "=MIN(CashFlow!D2:D" + str(last) + ")", MONEY, True),
        ("LTV (per customer)", f"=({AS('arpa')}*{AS('gross_margin')})/{AS('churn')}", MONEY2, False),
        ("LTV : CAC", f"=(({AS('arpa')}*{AS('gross_margin')})/{AS('churn')})/{AS('cac')}", '0.0"x"', True),
        ("CAC payback (months)", f"={AS('cac')}/({AS('arpa')}*{AS('gross_margin')})", '0.0', False),
    ]
    _hdr(sm, 3, ["Metric", "Value"])
    for i, (label, val, fmt, emphasize) in enumerate(metrics):
        r = 4 + i
        lc = sm.cell(row=r, column=1, value=label)
        vc = sm.cell(row=r, column=2, value=val)
        vc.number_format = fmt
        vc.alignment = Alignment(horizontal="right")
        lc.border = vc.border = BORDER
        if emphasize:
            lc.font = vc.font = Font(bold=True, color=PRIMARY_HEX)
    sm.column_dimensions["A"].width = 34
    sm.column_dimensions["B"].width = 18
    sm.cell(row=4 + len(metrics) + 1, column=1,
            value=f"Generated by Zilo · {datetime.utcnow().strftime('%d %b %Y')}").font = MUTED

    # ── Sheet 6: Charts (native Excel charts, recalc with the model) ───────────
    ch = wb.create_sheet("Charts")
    ch.sheet_view.showGridLines = False
    cats = Reference(rv, min_col=1, min_row=2, max_row=last)  # month numbers

    def _line(title, src_ws, col, anchor, color=None):
        c = LineChart()
        c.title = title
        c.height, c.width = 7.5, 15
        c.style = 12
        data = Reference(src_ws, min_col=col, min_row=1, max_row=last)  # incl. header = series name
        c.add_data(data, titles_from_data=True)
        c.set_categories(cats)
        c.x_axis.title = "Month"
        c.y_axis.numFmt = "#,##0"
        if c.series and color:
            c.series[0].graphicalProperties.line.solidFill = color
        ch.add_chart(c, anchor)

    _line("MRR", rv, 6, "B2", PRIMARY_HEX)
    _line("Customers", rv, 5, "L2", "10B981")
    _line("Ending Cash", cf, 4, "B18", "F59E0B")

    # EBITDA bar
    bar = BarChart()
    bar.title = "Monthly EBITDA"
    bar.height, bar.width = 7.5, 15
    bar.style = 10
    bar.add_data(Reference(pl, min_col=7, min_row=1, max_row=last), titles_from_data=True)
    bar.set_categories(cats)
    bar.x_axis.title = "Month"
    bar.y_axis.numFmt = "#,##0"
    ch.add_chart(bar, "L18")

    # ── Cover Sheet (First Tab) ──────────────────────────────────────────────
    cv = wb.create_sheet("Cover", 0)
    cv.sheet_view.showGridLines = False

    # Accent Strip (Column A)
    BRAND_ACCENT_FILL = PatternFill("solid", fgColor=PRIMARY_HEX)
    for r_idx in range(1, 30):
        cv.cell(row=r_idx, column=1).fill = BRAND_ACCENT_FILL

    cv.column_dimensions["A"].width = 3
    cv.column_dimensions["B"].width = 4
    cv.column_dimensions["C"].width = 45

    # Title
    c_title = cv.cell(row=4, column=3, value=f"{business_name or 'Zilo'} SaaS Financial Model")
    c_title.font = Font(name="Arial", bold=True, size=22, color=PRIMARY_HEX)

    # Subtitle
    c_sub = cv.cell(row=5, column=3, value="FINANCIAL PROJECTIONS & SCENARIO ANALYSIS")
    c_sub.font = Font(name="Arial", bold=True, size=11, color="6B7280")

    # Divider line
    BORDER_SIDE = Side(style="medium", color="D1D5DB")
    for col_idx in range(3, 8):
        cv.cell(row=6, column=col_idx).border = Border(bottom=BORDER_SIDE)

    # Metadata Block
    cv.cell(row=8, column=3, value="Prepared For:").font = Font(name="Arial", italic=True, size=9, color="6B7280")
    cv.cell(row=9, column=3, value=business_name or "Internal Use").font = Font(name="Arial", bold=True, size=11, color="111827")

    cv.cell(row=11, column=3, value="Date:").font = Font(name="Arial", italic=True, size=9, color="6B7280")
    cv.cell(row=12, column=3, value=datetime.utcnow().strftime('%B %d, %Y')).font = Font(name="Arial", bold=True, size=11, color="111827")

    cv.cell(row=14, column=3, value="Confidentiality:").font = Font(name="Arial", italic=True, size=9, color="6B7280")
    cv.cell(row=15, column=3, value="CONFIDENTIAL — FOR INTERNAL USE ONLY").font = Font(name="Arial", bold=True, size=10, color="991B1B")

    # Open on the Cover sheet first.
    wb.active = wb.sheetnames.index("Cover")

    # Turn off grid lines and set tab colors on all sheets
    for ws in wb.worksheets:
        ws.sheet_view.showGridLines = False
        name = ws.title.lower()
        if "cover" in name:
            ws.sheet_properties.tabColor = PRIMARY_HEX
        elif "assumption" in name:
            ws.sheet_properties.tabColor = "22C55E"  # Green
        elif "revenue" in name or "mrr" in name:
            ws.sheet_properties.tabColor = "111111"  # Dark
        elif "p&l" in name or "cashflow" in name or "opex" in name or "burn" in name or "cost" in name:
            ws.sheet_properties.tabColor = "EF4444"  # Red
        elif "summary" in name or "kpi" in name:
            ws.sheet_properties.tabColor = "F59E0B"  # Amber
        elif "chart" in name or "analysis" in name:
            ws.sheet_properties.tabColor = "8B5CF6"  # Purple

    wb.save(str(filepath))
    return str(filepath)


# ─────────────────────────────────────────────────────────────────────────────
# Spreadsheet styling and utility helpers (contrast and tinting)
# ─────────────────────────────────────────────────────────────────────────────

def _get_tint_color(hex_color: str, factor: float) -> str:
    """Mix the hex color with white. factor = 0.0 is pure white, 1.0 is original color."""
    hex_color = hex_color.lstrip("#")
    if len(hex_color) == 3:
        hex_color = "".join(c*2 for c in hex_color)
    if len(hex_color) != 6:
        return "FFFFFF"
    try:
        r = int(hex_color[0:2], 16)
        g = int(hex_color[2:4], 16)
        b = int(hex_color[4:6], 16)
    except ValueError:
        return "FFFFFF"
    
    # Mix with white (255, 255, 255)
    r = int(r * factor + 255 * (1 - factor))
    g = int(g * factor + 255 * (1 - factor))
    b = int(b * factor + 255 * (1 - factor))
    
    return f"{r:02X}{g:02X}{b:02X}"


def _is_light_color(hex_color: str) -> bool:
    """Check if hex color is light (luminance > 180)."""
    hex_color = hex_color.lstrip("#")
    if len(hex_color) == 3:
        hex_color = "".join(c*2 for c in hex_color)
    if len(hex_color) != 6:
        return True
    try:
        r = int(hex_color[0:2], 16)
        g = int(hex_color[2:4], 16)
        b = int(hex_color[4:6], 16)
    except ValueError:
        return True
    # Perceptive luminance formula
    luminance = 0.299 * r + 0.587 * g + 0.114 * b
    return luminance > 180


def _write_xlsx_table_rows(
    ws,
    columns: list[dict],
    rows: list[dict],
    start_row: int,
    style: dict,
    total_keys: set,
    hl: tuple | None,
    PRIMARY_HEX: str,
):
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
    
    ALT_ROW_HEX = _get_tint_color(PRIMARY_HEX, 0.03)  # very subtle 3% brand tint
    TOTAL_ROW_HEX = _get_tint_color(PRIMARY_HEX, 0.08)  # 8% brand tint for totals
    
    HEADER_TEXT_COLOR = "000000" if _is_light_color(PRIMARY_HEX) else "FFFFFF"
    
    THIN = Side(style="thin", color="D1D5DB")
    BORDER = Border(left=THIN, right=THIN, top=THIN, bottom=THIN)
    FMT = {
        "money": "$#,##0;($#,##0);-",
        "int": "#,##0;(#,##0);-",
        "pct": "0.0%",
        "date": "yyyy-mm-dd hh:mm"
    }
    
    BAND_FILL = PatternFill("solid", fgColor=ALT_ROW_HEX)
    SECTION_FILL = PatternFill("solid", fgColor=PRIMARY_HEX)
    TOTAL_FILL = PatternFill("solid", fgColor=TOTAL_ROW_HEX)
    FLAG_FILL = PatternFill("solid", fgColor="FEF3C7")  # Amber row highlight
    
    # Financial Highlight Row Fills & Fonts
    GP_FILL = PatternFill("solid", fgColor="E6F4EA")  # Soft green
    GP_FONT = Font(name="Arial", bold=True, color="065F46", size=10)
    
    OPEX_FILL = PatternFill("solid", fgColor="FCE8E6")  # Soft red
    OPEX_FONT = Font(name="Arial", bold=True, color="991B1B", size=10)
    
    NP_FILL = PatternFill("solid", fgColor="FEF7E0")  # Soft yellow
    NP_FONT = Font(name="Arial", bold=True, color="92400E", size=10)
    
    for i, rowdata in enumerate(rows):
        r = start_row + i
        
        # A. Detect Row Types
        val_first = rowdata.get(columns[0]["key"])
        val_first_str = str(val_first).strip() if val_first is not None else ""
        
        # Check if other columns are empty or have minor divider indicators
        other_non_empty_count = 0
        for col in columns[1:]:
            v = rowdata.get(col["key"])
            if v is not None and str(v).strip() != "":
                other_non_empty_count += 1
        
        is_section = False
        if val_first_str:
            clean_text = re.sub(r"\*\*|\*", "", val_first_str).strip()
            is_upper = clean_text.isupper()
            starts_with_num = re.match(r"^(?:Section|Part|Category|Phase|\d+|[IVXLCDM]+)[\s.:\-]", clean_text, re.IGNORECASE) is not None
            if other_non_empty_count == 0:
                is_section = True
            elif (is_upper or "section" in clean_text.lower() or "break" in clean_text.lower() or starts_with_num) and other_non_empty_count <= 1:
                is_section = True
        
        # Set Row heights dynamically per type
        if is_section:
            ws.row_dimensions[r].height = 22
        else:
            ws.row_dimensions[r].height = 17
            
        # B. Check Financial Highlights
        header_clean = re.sub(r"\*\*|\*", "", val_first_str).lower().strip()
        is_gp = "gross profit" in header_clean or "gross margin" in header_clean or header_clean == "gp"
        is_opex = "operating expense" in header_clean or "opex" in header_clean or "total expense" in header_clean or "operating cost" in header_clean
        is_np = "net profit" in header_clean or "net income" in header_clean or "net margin" in header_clean or "net earnings" in header_clean
        is_any_total = "total" in header_clean and not (is_gp or is_opex or is_np)
        
        # C. Process Highlight Conditions
        flag = False
        if hl:
            hk, op, hv = hl
            v = rowdata.get(hk)
            try:
                flag = (op == ">" and v > hv) or (op == "<" and v < hv) or (op == "==" and v == hv)
            except TypeError:
                flag = False
                
        # D. Write and style cells
        for j, col in enumerate(columns, start=1):
            val = rowdata.get(col["key"])
            cell = ws.cell(row=r, column=j, value=val)
            cell.border = BORDER
            
            # Formats
            f = col.get("fmt")
            if f in ("money", "int", "pct"):
                cell.number_format = FMT[f]
            elif f == "date":
                cell.number_format = FMT[f]
                
            # Alignments
            col_header = str(col.get("header") or "").lower()
            val_str = str(val).strip() if val is not None else ""
            
            if f in ("money", "int", "pct") or val_str.startswith("="):
                align_h = "right"
            elif f == "date" or col_header in ("status", "flag", "code", "id") or len(val_str) <= 5:
                align_h = "center"
            else:
                align_h = "left"
                
            cell.alignment = Alignment(horizontal=align_h, vertical="center")
            
            # Fills and Fonts
            if is_section:
                cell.fill = SECTION_FILL
                cell.font = Font(name="Arial", bold=True, color=HEADER_TEXT_COLOR, size=10)
                cell.alignment = Alignment(horizontal="left", vertical="center")
            elif is_gp:
                cell.fill = GP_FILL
                cell.font = GP_FONT
            elif is_opex:
                cell.fill = OPEX_FILL
                cell.font = OPEX_FONT
            elif is_np:
                cell.fill = NP_FILL
                cell.font = NP_FONT
            elif is_any_total:
                cell.fill = TOTAL_FILL
                cell.font = Font(name="Arial", bold=True, color="111827", size=10)
            else:
                if flag:
                    cell.fill = FLAG_FILL
                elif i % 2 == 1:
                    cell.fill = BAND_FILL
                    
                # Content-based font color coding
                is_note_col = any(keyword in col_header for keyword in ("note", "comment", "remark", "desc", "unit", "spec"))
                
                if is_note_col:
                    cell.font = Font(name="Arial", italic=True, color="6B7280", size=10)
                elif val_str.startswith("="):
                    if "!" in val_str:
                        cell.font = Font(name="Arial", bold=True, color="15803D", size=10)
                    else:
                        cell.font = Font(name="Arial", color="111827", size=10)
                elif isinstance(val, (int, float)):
                    cell.font = Font(name="Arial", color="1A56DB", size=10)
                else:
                    cell.font = Font(name="Arial", color="111827", size=10)


# ─────────────────────────────────────────────────────────────────────────────
# Generic CRM data export (.xlsx) — branded table with totals + highlights
# ─────────────────────────────────────────────────────────────────────────────
#
# columns: list of dicts — {"header", "key", "fmt"?, "width"?, "align"?}
#   fmt: "money" | "int" | "pct" | "date" | None (text)
# rows:    list of dicts keyed by the column "key"s
# total_keys: keys to SUM in a footer row
# highlight: optional (key, operator, value) — shades a row's cells when the row's
#            value for `key` satisfies the comparison (e.g. ("remaining", ">", 0))
# ─────────────────────────────────────────────────────────────────────────────

def generate_data_xlsx(
    *,
    title: str,
    columns: list[dict],
    rows: list[dict],
    business_name: str = "",
    style: dict | None = None,
    total_keys: set | None = None,
    sheet_name: str = "Data",
    filename: str | None = None,
) -> str:
    """Build a branded, filterable Excel table from a list of row dicts. Returns path."""
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side

    style = style or {}
    total_keys = total_keys or set()

    filename = filename or f"zilo_export_{uuid.uuid4().hex[:8]}.xlsx"
    if not filename.endswith(".xlsx"):
        filename += ".xlsx"
    filepath = TEMP_DIR / filename

    PRIMARY_HEX = _safe_hex(style.get("primary_color", ""), "#4F46E5").lstrip("#").upper()
    HEADER_FILL = PatternFill("solid", fgColor=PRIMARY_HEX)
    HEADER_TEXT_COLOR = "000000" if _is_light_color(PRIMARY_HEX) else "FFFFFF"
    HEADER_FONT = Font(name="Arial", color=HEADER_TEXT_COLOR, bold=True, size=10)
    WHITE = Font(color="FFFFFF", bold=True)
    BOLD = Font(bold=True)
    THIN = Side(style="thin", color="D1D5DB")
    BORDER = Border(left=THIN, right=THIN, top=THIN, bottom=THIN)
    CENTER = Alignment(horizontal="center", vertical="center")
    FMT = {
        "money": "$#,##0;($#,##0);-",
        "int": "#,##0;(#,##0);-",
        "pct": "0.0%",
        "date": "yyyy-mm-dd hh:mm"
    }

    wb = Workbook()
    ws = wb.active
    ws.title = sheet_name[:31] or "Data"
    ws.sheet_view.showGridLines = False

    # Row 1: Header Band (Company Brand Color)
    ws.row_dimensions[1].height = 44
    for col_idx in range(1, max(20, len(columns) + 1)):
        ws.cell(row=1, column=col_idx).fill = HEADER_FILL
    ws.cell(row=1, column=1, value=f"{business_name + ' — ' if business_name else ''}{title}").font = Font(name="Arial", bold=True, size=16, color=HEADER_TEXT_COLOR)
    ws.cell(row=1, column=1).alignment = Alignment(horizontal="left", vertical="center")

    # Row 2: Subtitle
    ws.row_dimensions[2].height = 20
    ws.cell(row=2, column=1, value=f"{len(rows)} rows · generated {datetime.utcnow().strftime('%d %b %Y %H:%M')} UTC").font = Font(name="Arial", italic=True, size=9, color="6B7280")
    ws.cell(row=2, column=1).alignment = Alignment(horizontal="left", vertical="center")

    head_row = 4
    ws.row_dimensions[head_row].height = 24
    for j, col in enumerate(columns, start=1):
        c = ws.cell(row=head_row, column=j, value=col["header"])
        c.fill = HEADER_FILL
        c.font = HEADER_FONT
        c.alignment = CENTER
        c.border = BORDER

    hl = style.get("_highlight")  # (key, op, value)
    _write_xlsx_table_rows(
        ws,
        columns=columns,
        rows=rows,
        start_row=head_row + 1,
        style=style,
        total_keys=total_keys,
        hl=hl,
        PRIMARY_HEX=PRIMARY_HEX,
    )

    # Totals footer
    if total_keys and rows:
        r = head_row + 1 + len(rows)
        ws.row_dimensions[r].height = 22
        
        TOTAL_ROW_HEX = _get_tint_color(PRIMARY_HEX, 0.08)
        TOTAL_FILL = PatternFill("solid", fgColor=TOTAL_ROW_HEX)
        
        DOUBLE = Side(style="double", color="111827")
        TOTALS_BORDER = Border(top=THIN, bottom=DOUBLE, left=THIN, right=THIN)
        
        for j, col in enumerate(columns, start=1):
            cell = ws.cell(row=r, column=j)
            cell.border = TOTALS_BORDER
            cell.fill = TOTAL_FILL
            
            if j == 1:
                cell.value = "TOTAL"
                cell.font = Font(name="Arial", bold=True, color="111827", size=10)
                cell.alignment = Alignment(horizontal="left", vertical="center")
            elif col["key"] in total_keys:
                first = head_row + 1
                last = head_row + len(rows)
                colletter = cell.column_letter
                cell.value = f"=SUM({colletter}{first}:{colletter}{last})"
                cell.font = Font(name="Arial", bold=True, color="111827", size=10)
                cell.number_format = FMT.get(col.get("fmt"), "#,##0;(#,##0);-")
                cell.alignment = Alignment(horizontal="right", vertical="center")

    # Column widths (dynamic) + freeze + autofilter
    for j, col in enumerate(columns, start=1):
        # Calculate max length of values in this column
        max_len = 0
        for row in rows:
            val_str = str(row.get(col["key"]) or "")
            if len(val_str) > max_len:
                max_len = len(val_str)
        col_width = max(len(col["header"]), max_len) + 4
        col_letter = ws.cell(row=head_row, column=j).column_letter
        ws.column_dimensions[col_letter].width = max(12, min(col_width, 50))

    ws.freeze_panes = "A5"
    if columns:
        last_col = ws.cell(row=head_row, column=len(columns)).column_letter
        ws.auto_filter.ref = f"A{head_row}:{last_col}{head_row + len(rows)}"

    wb.save(str(filepath))
    return str(filepath)


# ─────────────────────────────────────────────────────────────────────────────
# Multi-sheet workbook (e.g. weekly report) — one table per sheet + totals
# ─────────────────────────────────────────────────────────────────────────────
#
# sheets: list of dicts — {"title", "columns", "rows", "total_keys"?, "highlight"?}
#   (same column/row shape as generate_data_xlsx)
# An optional leading "Summary" sheet of key→value pairs can be supplied via
# `summary` = list of (label, value, fmt).
# ─────────────────────────────────────────────────────────────────────────────

def generate_multi_sheet_xlsx(
    *,
    title: str,
    sheets: list[dict],
    summary: list[tuple] | None = None,
    business_name: str = "",
    style: dict | None = None,
    filename: str | None = None,
) -> str:
    """Build a multi-tab workbook (one table per sheet) with an optional Summary tab."""
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side

    style = style or {}
    filename = filename or f"zilo_report_{uuid.uuid4().hex[:8]}.xlsx"
    if not filename.endswith(".xlsx"):
        filename += ".xlsx"
    filepath = TEMP_DIR / filename

    PRIMARY_HEX = _safe_hex(style.get("primary_color", ""), "#4F46E5").lstrip("#").upper()
    HEADER_FILL = PatternFill("solid", fgColor=PRIMARY_HEX)
    HEADER_TEXT_COLOR = "000000" if _is_light_color(PRIMARY_HEX) else "FFFFFF"
    HEADER_FONT = Font(name="Arial", color=HEADER_TEXT_COLOR, bold=True, size=10)
    WHITE = Font(color="FFFFFF", bold=True)
    BOLD = Font(bold=True)
    THIN = Side(style="thin", color="D1D5DB")
    BORDER = Border(left=THIN, right=THIN, top=THIN, bottom=THIN)
    CENTER = Alignment(horizontal="center", vertical="center")
    RIGHT = Alignment(horizontal="right")
    FMT = {
        "money": "$#,##0;($#,##0);-",
        "int": "#,##0;(#,##0);-",
        "pct": "0.0%",
        "date": "yyyy-mm-dd hh:mm"
    }

    wb = Workbook()

    def _write_table(ws, tbl):
        cols = tbl["columns"]
        rows = tbl.get("rows") or []
        total_keys = tbl.get("total_keys") or set()
        hl = tbl.get("highlight")  # (key, op, value)
        head = 4

        # Turn off gridlines
        ws.sheet_view.showGridLines = False

        # Row 1: Header Band (Company Brand Color)
        ws.row_dimensions[1].height = 44
        for col_idx in range(1, max(20, len(cols) + 1)):
            ws.cell(row=1, column=col_idx).fill = HEADER_FILL
        title_val = tbl.get("title", ws.title)
        ws.cell(row=1, column=1, value=title_val).font = Font(name="Arial", bold=True, size=16, color=HEADER_TEXT_COLOR)
        ws.cell(row=1, column=1).alignment = Alignment(horizontal="left", vertical="center")

        # Row 2: Subtitle
        ws.row_dimensions[2].height = 20
        subtitle_text = f"{business_name} · {len(rows)} rows" if business_name else f"{len(rows)} rows"
        ws.cell(row=2, column=1, value=subtitle_text).font = Font(name="Arial", italic=True, size=9, color="6B7280")
        ws.cell(row=2, column=1).alignment = Alignment(horizontal="left", vertical="center")

        # Row 4: Table headers
        ws.row_dimensions[head].height = 24
        for j, col in enumerate(cols, start=1):
            c = ws.cell(row=head, column=j, value=col["header"])
            c.fill, c.font, c.alignment, c.border = HEADER_FILL, HEADER_FONT, CENTER, BORDER

        _write_xlsx_table_rows(
            ws,
            columns=cols,
            rows=rows,
            start_row=head + 1,
            style=style,
            total_keys=total_keys,
            hl=hl,
            PRIMARY_HEX=PRIMARY_HEX,
        )

        if total_keys and rows:
            r = head + 1 + len(rows)
            ws.row_dimensions[r].height = 22
            
            TOTAL_ROW_HEX = _get_tint_color(PRIMARY_HEX, 0.08)
            TOTAL_FILL = PatternFill("solid", fgColor=TOTAL_ROW_HEX)
            
            DOUBLE = Side(style="double", color="111827")
            TOTALS_BORDER = Border(top=THIN, bottom=DOUBLE, left=THIN, right=THIN)
            
            for j, col in enumerate(cols, start=1):
                cell = ws.cell(row=r, column=j)
                cell.border = TOTALS_BORDER
                cell.fill = TOTAL_FILL
                
                if j == 1:
                    cell.value = "TOTAL"
                    cell.font = Font(name="Arial", bold=True, color="111827", size=10)
                    cell.alignment = Alignment(horizontal="left", vertical="center")
                elif col["key"] in total_keys:
                    first = head + 1
                    last = head + len(rows)
                    letter = cell.column_letter
                    cell.value = f"=SUM({letter}{first}:{letter}{last})"
                    cell.font = Font(name="Arial", bold=True, color="111827", size=10)
                    cell.number_format = FMT.get(col.get("fmt"), "#,##0;(#,##0);-")
                    cell.alignment = Alignment(horizontal="right", vertical="center")

        # Column widths (dynamic) + freeze + autofilter
        for j, col in enumerate(cols, start=1):
            max_len = 0
            for row in rows:
                val_str = str(row.get(col["key"]) or "")
                if len(val_str) > max_len:
                    max_len = len(val_str)
            col_width = max(len(col["header"]), max_len) + 4
            col_letter = ws.cell(row=head, column=j).column_letter
            ws.column_dimensions[col_letter].width = max(12, min(col_width, 50))

        ws.freeze_panes = "A5"
        if cols:
            last_col = ws.cell(row=head, column=len(cols)).column_letter
            ws.auto_filter.ref = f"A{head}:{last_col}{head + len(rows)}"

    # Summary tab first
    if summary is not None:
        sm = wb.active
        sm.title = "Summary"
        sm.sheet_view.showGridLines = False

        # Row 1: Header Band (Company Brand Color)
        sm.row_dimensions[1].height = 44
        for col_idx in range(1, 21):
            sm.cell(row=1, column=col_idx).fill = HEADER_FILL
        sm.cell(row=1, column=1, value=f"{business_name + ' — ' if business_name else ''}{title}").font = Font(name="Arial", bold=True, size=16, color=HEADER_TEXT_COLOR)
        sm.cell(row=1, column=1).alignment = Alignment(horizontal="left", vertical="center")

        # Row 2: Subtitle
        sm.row_dimensions[2].height = 20
        sm.cell(row=2, column=1, value=f"Generated {datetime.utcnow().strftime('%d %b %Y %H:%M')} UTC").font = Font(name="Arial", italic=True, size=9, color="6B7280")
        sm.cell(row=2, column=1).alignment = Alignment(horizontal="left", vertical="center")

        # Table headers
        sm.row_dimensions[4].height = 24
        for j, h in enumerate(("Metric", "Value"), start=1):
            c = sm.cell(row=4, column=j, value=h)
            c.fill, c.font, c.alignment, c.border = HEADER_FILL, HEADER_FONT, CENTER, BORDER
            
        ALT_ROW_HEX = _get_tint_color(PRIMARY_HEX, 0.03)
        BAND_FILL = PatternFill("solid", fgColor=ALT_ROW_HEX)
        
        for i, (label, value, fmt) in enumerate(summary):
            r = 5 + i
            sm.row_dimensions[r].height = 20
            
            c1 = sm.cell(row=r, column=1, value=label)
            c1.border = BORDER
            c1.font = Font(name="Arial", bold=True, color="111827", size=10)
            c1.alignment = Alignment(horizontal="left", vertical="center")
            
            c2 = sm.cell(row=r, column=2, value=value)
            c2.border = BORDER
            c2.font = Font(name="Arial", color="111827", size=10)
            
            if fmt in FMT:
                c2.number_format = FMT[fmt]
                c2.alignment = Alignment(horizontal="right", vertical="center")
                if isinstance(value, (int, float)):
                    c2.font = Font(name="Arial", color="1A56DB", size=10)
            else:
                c2.alignment = Alignment(horizontal="left", vertical="center")
                
            if i % 2 == 1:
                c1.fill = BAND_FILL
                c2.fill = BAND_FILL
                
        sm.column_dimensions["A"].width = 32
        sm.column_dimensions["B"].width = 18
        first = True
    else:
        first = True

    for tbl in sheets:
        name = (tbl.get("title") or "Sheet")[:31]
        if first and summary is None:
            ws = wb.active
            ws.title = name
            first = False
        else:
            ws = wb.create_sheet(name)
        _write_table(ws, tbl)

    # Set Tab Colors on all sheets
    for ws in wb.worksheets:
        name = ws.title.lower()
        if "cover" in name:
            ws.sheet_properties.tabColor = PRIMARY_HEX
        elif "assumption" in name:
            ws.sheet_properties.tabColor = "22C55E"  # Green
        elif "revenue" in name or "mrr" in name or "sales" in name:
            ws.sheet_properties.tabColor = "111111"  # Dark
        elif "p&l" in name or "cashflow" in name or "opex" in name or "burn" in name or "cost" in name or "expense" in name:
            ws.sheet_properties.tabColor = "EF4444"  # Red
        elif "summary" in name or "kpi" in name:
            ws.sheet_properties.tabColor = "F59E0B"  # Amber
        elif "chart" in name or "analysis" in name:
            ws.sheet_properties.tabColor = "8B5CF6"  # Purple
        elif "tracker" in name or "pipeline" in name or "customer" in name:
            ws.sheet_properties.tabColor = "3B82F6"  # Blue

    wb.save(str(filepath))
    return str(filepath)


# ─────────────────────────────────────────────────────────────────────────────
# Excel reader — parse an uploaded .xlsx into a list of row dicts
# ─────────────────────────────────────────────────────────────────────────────

def read_xlsx(file_bytes: bytes, *, sheet: str | None = None, max_rows: int = 10000) -> list[dict]:
    """Parse the first (or named) worksheet into a list of dicts keyed by the
    normalized header row (lower-cased, stripped). Fully-blank rows are skipped.
    """
    import io as _io
    from openpyxl import load_workbook

    wb = load_workbook(_io.BytesIO(file_bytes), read_only=True, data_only=True)
    ws = wb[sheet] if sheet and sheet in wb.sheetnames else wb[wb.sheetnames[0]]

    rows_iter = ws.iter_rows(values_only=True)
    headers: list[str] = []
    for raw in rows_iter:
        if raw and any(c is not None and str(c).strip() for c in raw):
            headers = [str(c).strip().lower() if c is not None else "" for c in raw]
            break
    if not headers:
        return []

    out: list[dict] = []
    for raw in rows_iter:
        if raw is None or not any(c is not None and str(c).strip() for c in raw):
            continue
        row = {}
        for i, h in enumerate(headers):
            if h:
                row[h] = raw[i] if i < len(raw) else None
        out.append(row)
        if len(out) >= max_rows:
            break
    wb.close()
    return out


# ─────────────────────────────────────────────────────────────────────────────
# Cleanup helper
# ─────────────────────────────────────────────────────────────────────────────
def cleanup_file(filepath: str) -> None:
    try:
        os.remove(filepath)
    except Exception:
        pass
